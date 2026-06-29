from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "figures"
OUT_DOCX = ROOT / "bao_cao_tien_do_nghien_cuu_gvhd_2026-06-13.docx"
FIG_PATH = FIG_DIR / "kien_truc_retrieval_reasoning_2026-06-13.png"
CONTRIB_DIR = FIG_DIR / "contribution1_assets"


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    font_path = Path("C:/Windows/Fonts") / font_name
    return ImageFont.truetype(str(font_path), size=size)


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill):
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=4, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x = x0 + (x1 - x0 - tw) / 2
    y = y0 + (y1 - y0 - th) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=4, align="center")


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, fill, outline, title_fill=(20, 20, 20)):
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, title, get_font(24, bold=True), title_fill)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(60, 60, 60), width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        p1 = (ex - 18 * sign, ey - 8)
        p2 = (ex - 18 * sign, ey + 8)
    else:
        sign = 1 if dy >= 0 else -1
        p1 = (ex - 8, ey - 18 * sign)
        p2 = (ex + 8, ey - 18 * sign)
    draw.polygon([end, p1, p2], fill=color)


def build_diagram(out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2200, 1500), (250, 250, 252))
    draw = ImageDraw.Draw(img)

    title_font = get_font(40, bold=True)
    body_font = get_font(22)
    section_font = get_font(28, bold=True)

    draw.text((60, 40), "Kiến trúc mô hình: hiện trạng retrieval và hướng phát triển toàn trình", font=title_font, fill=(10, 40, 95))

    # Panel A
    panel_a = (50, 130, 2150, 700)
    panel_b = (50, 780, 2150, 1425)
    draw.rounded_rectangle(panel_a, radius=20, outline=(20, 82, 147), width=4, fill=(240, 247, 255))
    draw.rounded_rectangle(panel_b, radius=20, outline=(0, 110, 82), width=4, fill=(241, 252, 247))
    draw.text((80, 150), "A. Kiến trúc retrieval hiện tại theo GSR-CACL", font=section_font, fill=(20, 82, 147))
    draw.text((80, 800), "B. Kiến trúc đề xuất cho retrieval + reasoning toàn trình", font=section_font, fill=(0, 110, 82))

    a_boxes = [
        ((100, 260, 360, 390), "Câu hỏi +\nmetadata truy vấn", (255, 255, 255), (20, 82, 147)),
        ((430, 260, 710, 390), "FAISS lấy tập\nứng viên tài liệu", (255, 255, 255), (20, 82, 147)),
        ((780, 260, 1070, 390), "Trích bảng markdown\nđầu tiên", (255, 255, 255), (20, 82, 147)),
        ((1140, 210, 1450, 340), "Template matching\nvà xây Constraint KG", (255, 255, 255), (20, 82, 147)),
        ((1140, 400, 1450, 530), "Constraint score", (255, 255, 255), (20, 82, 147)),
        ((1520, 210, 1840, 340), "Edge-aware GAT\n→ kg_embed", (255, 255, 255), (20, 82, 147)),
        ((1520, 400, 1840, 530), "Entity score\n(company/year/sector)", (255, 255, 255), (20, 82, 147)),
        ((1900, 260, 2110, 390), "Joint scorer\n→ Top-k", (255, 255, 255), (20, 82, 147)),
    ]
    for box, text, fill, outline in a_boxes:
        draw_box(draw, box, text, fill, outline)

    draw_arrow(draw, (360, 325), (430, 325), color=(20, 82, 147))
    draw_arrow(draw, (710, 325), (780, 325), color=(20, 82, 147))
    draw_arrow(draw, (1070, 325), (1140, 275), color=(20, 82, 147))
    draw_arrow(draw, (1070, 325), (1140, 465), color=(20, 82, 147))
    draw_arrow(draw, (1450, 275), (1520, 275), color=(20, 82, 147))
    draw_arrow(draw, (1450, 465), (1900, 355), color=(20, 82, 147))
    draw_arrow(draw, (1840, 275), (1900, 325), color=(20, 82, 147))
    draw_arrow(draw, (1840, 465), (1900, 355), color=(20, 82, 147))

    draw.text(
        (120, 565),
        "Điểm nghẽn chính của kiến trúc hiện tại: retrieval mới mạnh ở cấp tài liệu, KG vẫn thiên về cấu trúc reranking\n"
        "và chưa đóng vai trò đầy đủ trong việc định vị toán hạng, lập phương trình hay kiểm chứng suy luận số học.",
        font=body_font,
        fill=(40, 40, 40),
        spacing=6,
    )

    # Panel B
    b_boxes = [
        ((95, 915, 315, 1040), "Câu hỏi", (255, 255, 255), (0, 110, 82)),
        ((360, 860, 620, 995), "Phân tích truy vấn:\ncompany / time /\nmetric / operation", (255, 255, 255), (0, 110, 82)),
        ((360, 1035, 620, 1170), "Ontology metadata:\nstatement, quarter,\nunit, scale, aliases", (255, 255, 255), (0, 110, 82)),
        ((680, 915, 960, 1040), "Retrieval phân cấp:\ndocument → table/\nsection → atom", (255, 255, 255), (0, 110, 82)),
        ((1015, 835, 1320, 970), "Financial Evidence Graph:\ncell, text, footnote,\nequation, provenance", (255, 255, 255), (0, 110, 82)),
        ((1015, 1080, 1320, 1215), "Evidence triage:\nlọc top-3 nhiễu,\nneo đúng toán hạng", (255, 255, 255), (0, 110, 82)),
        ((1380, 835, 1685, 970), "Reasoning planner:\nchọn phép toán,\nsinh DSL/Python", (255, 255, 255), (0, 110, 82)),
        ((1380, 1080, 1685, 1215), "Executor:\nthực thi số học", (255, 255, 255), (0, 110, 82)),
        ((1750, 915, 2045, 1040), "Verifier:\nkiểm tra grounding,\nunit, year,\nconstraint", (255, 255, 255), (0, 110, 82)),
    ]
    for box, text, fill, outline in b_boxes:
        draw_box(draw, box, text, fill, outline)

    draw_arrow(draw, (315, 975), (360, 925), color=(0, 110, 82))
    draw_arrow(draw, (315, 975), (360, 1100), color=(0, 110, 82))
    draw_arrow(draw, (620, 925), (680, 965), color=(0, 110, 82))
    draw_arrow(draw, (620, 1100), (680, 990), color=(0, 110, 82))
    draw_arrow(draw, (960, 975), (1015, 900), color=(0, 110, 82))
    draw_arrow(draw, (960, 975), (1015, 1148), color=(0, 110, 82))
    draw_arrow(draw, (1320, 900), (1380, 900), color=(0, 110, 82))
    draw_arrow(draw, (1320, 1148), (1380, 1148), color=(0, 110, 82))
    draw_arrow(draw, (1685, 900), (1750, 960), color=(0, 110, 82))
    draw_arrow(draw, (1685, 1148), (1750, 995), color=(0, 110, 82))

    draw.text(
        (120, 1260),
        "Ý tưởng chốt: biến KG thành Financial Evidence Graph dùng chung cho retrieval và reasoning, để mô hình\n"
        "không chỉ chọn đúng tài liệu mà còn chọn đúng toán hạng, thực thi đúng chương trình và tự kiểm chứng kết quả.",
        font=body_font,
        fill=(35, 45, 35),
        spacing=6,
    )

    img.save(out_path)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, RGBColor(18, 60, 110)),
        ("Heading 1", 16, RGBColor(18, 60, 110)),
        ("Heading 2", 13, RGBColor(32, 32, 32)),
        ("Heading 3", 12, RGBColor(64, 64, 64)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Caption" not in doc.styles:
        caption_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        caption_style.font.size = Pt(10.5)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(90, 90, 90)
    else:
        caption_style = doc.styles["Caption"]
        caption_style.font.name = "Times New Roman"
        caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        caption_style.font.size = Pt(10.5)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(90, 90, 90)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("BÁO CÁO TIẾN ĐỘ RETRIEVAL VÀ SUY LUẬN TÀI CHÍNH")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer.add_run("Trang ")
    r1.font.name = "Times New Roman"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r1.font.size = Pt(9)
    add_page_number(footer)


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)


def add_display_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_captioned_picture(doc: Document, image_path: Path, caption: str, width: Inches | None = None) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=width or Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_reference_list(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.add_run(f"[{idx}] {item}")


def build_report(docx_path: Path, fig_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    configure_styles(doc)
    add_header_footer(doc)

    core = doc.core_properties
    core.title = "Báo cáo chiến lược truy xuất và suy luận tài chính dựa trên KG"
    core.author = "OpenAI Codex"
    core.subject = "Báo cáo nghiên cứu gửi GVHD"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.add_run("BÁO CÁO CHIẾN LƯỢC\nTRUY XUẤT VÀ SUY LUẬN TÀI CHÍNH DỰA TRÊN KG").bold = True
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(18, 60, 110)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(20)
    sub.add_run("Tổng hợp benchmark, hiện trạng triển khai, phân tích kỹ thuật, định hướng nghiên cứu\nvà kế hoạch phát triển toàn trình cho bài toán tài chính").font.size = Pt(13)
    for run in sub.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(140)
    info.add_run("Ngày cập nhật: 13/06/2026\nTài liệu phục vụ báo cáo với giáo viên hướng dẫn").font.size = Pt(12)
    for run in info.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_page_break()

    doc.add_heading("1. Tóm tắt điều hành", level=1)
    add_paragraphs(
        doc,
        [
            "Mục tiêu của báo cáo này là tổng hợp một cách hệ thống toàn bộ những gì đã được xây dựng trong repo hiện tại, đối chiếu với bài toán nghiên cứu thực sự cần giải quyết, và từ đó chốt lại một hướng triển khai đủ mạnh cho giai đoạn tiếp theo. Trọng tâm của đề tài không còn nên được hiểu là bài toán tăng điểm retrieval đơn thuần, mà là bài toán truy xuất đúng bằng chứng, định vị đúng toán hạng và thực hiện suy luận số học đáng tin cậy trong bối cảnh tài liệu tài chính dài, nhiều bảng, nhiều kỳ báo cáo và có nhiễu retrieval.",
            "Qua quá trình khảo sát benchmark, đọc tài liệu liên quan, phân tích nội dung trong contribution1.pdf và đối chiếu trực tiếp với code của hệ thống GSR-CACL hiện tại, có thể khẳng định rằng repo đã có một nền kỹ thuật tốt ở mức retrieval có cấu trúc. Tuy nhiên, hệ thống hiện tại vẫn chủ yếu là một prototype retrieval-centric, trong đó KG mới đóng vai trò hỗ trợ reranking tài liệu chứ chưa thực sự trở thành hạ tầng bằng chứng dùng chung cho cả retrieval và reasoning.",
            "Kết luận chiến lược quan trọng nhất của báo cáo là cần chuyển trọng tâm nghiên cứu từ Constraint KG cho retrieval sang Financial Evidence Graph dùng chung cho retrieval và reasoning. Trong hướng mới này, metadata không còn là một bonus yếu cho reranker, mà phải được nâng thành ontology truy xuất có kiểu; còn reasoning không nên chạy trực tiếp trên top-k context, mà phải vận hành trên một đồ thị bằng chứng cục bộ đã được triage, grounding và kiểm chứng.",
        ],
    )

    doc.add_heading("2. Bối cảnh, benchmark và bản chất bài toán", level=1)
    add_paragraphs(
        doc,
        [
            "Các hệ thống hỏi đáp tài chính hiện đại gặp hai khó khăn bản chất. Thứ nhất, thông tin không nằm trong một đoạn văn ngắn mà trải dài trên nhiều báo cáo, nhiều bảng và nhiều footnote. Thứ hai, đáp án cuối cùng thường không phải là một fact có sẵn mà là kết quả của một quá trình định vị toán hạng, chuẩn hóa đơn vị và thực hiện phép tính nhiều bước. Chính vì vậy, các thiết lập QA giả định context vàng có sẵn thường không phản ánh đúng độ khó thực sự của bài toán.",
            "Trong bối cảnh đó, T2-RAGBench là benchmark phù hợp nhất với hướng đề tài vì nó đặt retrieval thành yêu cầu bắt buộc trước reasoning. Thay vì đánh giá trực tiếp mô hình sinh đáp án trên context đã biết đúng, benchmark này buộc hệ thống phải truy xuất từ corpus tài chính hỗn hợp gồm text và bảng rồi mới reasoning. Đây là setting gần thực tế hơn và cũng là lý do vì sao repo hiện tại bám vào T2-RAGBench là một lựa chọn hợp lý.",
            "Tuy nhiên, đánh giá chỉ bằng MRR@3 hay Recall@k là chưa đủ. Trong bài toán tài chính, hoàn toàn có thể xảy ra tình huống top-3 có một context đúng nhưng hai context còn lại gây nhiễu mạnh ở mức toán hạng. Khi đó, retrieval nhìn có vẻ tốt nhưng reasoning vẫn thất bại. Do đó, benchmark cần được nâng từ đánh giá tài liệu sang đánh giá nhiều tầng, bao gồm retrieval, evidence grounding, reasoning và robustness dưới nhiễu.",
        ],
    )

    doc.add_heading("2.1. Vai trò của các bộ dữ liệu trong lộ trình nghiên cứu", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Các bộ dữ liệu hiện có không đóng vai trò trùng lặp, mà mỗi bộ đóng góp một góc nhìn khác nhau vào bài toán. Bảng dưới đây tóm tắt vai trò của từng benchmark trong lộ trình đề tài.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Bộ dữ liệu", "Vai trò chính", "Đặc trưng nổi bật", "Mục tiêu sử dụng trong đề tài"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9EAF7")
    set_repeat_table_header(table.rows[0])

    dataset_rows = [
        ("FinQA", "Suy luận tài chính có chương trình vàng", "Nhiều câu hỏi cần phép tính nhiều bước trên bảng và văn bản", "Huấn luyện và đánh giá executor/program reasoning"),
        ("ConvFinQA", "Suy luận hội thoại tài chính", "Phụ thuộc liên lượt hỏi, reasoning dài hơn", "Đánh giá khả năng bám ngữ cảnh và reasoning nhiều bước"),
        ("TAT-QA", "Hybrid table-text QA", "Kết hợp bằng chứng bảng và đoạn văn", "Đánh giá grounding liên phương thức"),
        ("DocFinQA", "Ngữ cảnh dài", "Tăng độ khó retrieval thực tế", "Stress-test retrieval và evidence triage"),
        ("T2-RAGBench", "Benchmark retrieval-first", "Không giả định oracle context", "Đánh giá retrieval và pipeline toàn trình"),
    ]
    for row in dataset_rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=10)

    doc.add_heading("3. Hiện trạng triển khai trong repo", level=1)
    add_paragraphs(
        doc,
        [
            "Qua phân tích mã nguồn, có thể xác định repo hiện tại gồm hai lớp chính. Lớp thứ nhất là các baseline retrieval trong thư mục baseline, bao gồm dense FAISS, hybrid BM25 cộng dense, HyDE và một số biến thể query rewriting. Lớp thứ hai là prototype đề xuất trong thư mục ours, nơi triển khai GSR-CACL với tham vọng đưa cấu trúc tài chính vào retrieval.",
            "Điều đáng lưu ý là contribution1.pdf không chỉ là phần trình bày ý tưởng, mà nhiều thành phần trong đó thực sự đã được hiện thực hóa ở mức prototype. Cụ thể, phần Graph-Structured Retrieval được phản ánh trong mô-đun xây KG theo template, Edge-aware GAT, constraint scoring và joint scorer; còn phần Constraint-Aware Contrastive Learning được phản ánh trong bộ sinh CHAP negatives, ba giai đoạn huấn luyện và loss kết hợp triplet với penalty vi phạm ràng buộc.",
            "Tuy vậy, việc hệ thống đã được cài đặt không đồng nghĩa rằng nó đã khép kín về mặt thực nghiệm. Qua đối chiếu code, có những điểm rất quan trọng: benchmark hiện tại phản ánh GSR inference rõ hơn là full GSR+CACL; constraint score hiện tại mới là tín hiệu cấu trúc xấp xỉ; và train/inference còn tồn tại lệch trong cách xây KG. Những phát hiện này rất quan trọng để tránh ngộ nhận về độ hoàn thiện của hệ thống.",
        ],
    )

    doc.add_heading("3.1. Bản đồ thành phần kỹ thuật hiện có", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Bảng sau tóm tắt các mô-đun quan trọng trong repo hiện tại, mục tiêu của từng mô-đun và đánh giá ngắn gọn về trạng thái sử dụng.")

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["Thành phần", "Đường dẫn chính", "Chức năng", "Đánh giá hiện trạng"]
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9EAD3")
    set_repeat_table_header(table2.rows[0])

    module_rows = [
        ("Dataset wrapper", "ours/source/src/gsr_cacl/datasets/wrappers.py", "Load T2-RAGBench, build corpus, query và metadata", "Đã dùng được, là nền tốt cho giai đoạn tiếp theo"),
        ("GSR retrieval", "ours/source/src/gsr_cacl/methods/gsr_retrieval.py", "FAISS candidate retrieval và rerank bằng joint scorer", "Đã có prototype thật sự ở mức retrieval"),
        ("KG builder", "ours/source/src/gsr_cacl/kg/builder.py", "Sinh Constraint KG từ bảng markdown theo template", "Hữu ích cho retrieval nhưng chưa đủ trung thành cho reasoning"),
        ("Template library", "ours/source/src/gsr_cacl/templates/library.py", "Thư viện template IFRS/GAAP", "Có thể tái sử dụng như lớp tri thức nền"),
        ("Edge-aware GAT", "ours/source/src/gsr_cacl/encoders/gat_encoder.py", "Mã hóa KG thành graph embedding", "Có giá trị ở mức structural reranking"),
        ("Constraint scoring", "ours/source/src/gsr_cacl/scoring/constraint_score.py", "Tạo điểm tuân thủ ràng buộc", "Cần thay bằng biểu diễn phương trình trung thành hơn"),
        ("Joint scorer", "ours/source/src/gsr_cacl/scoring/joint_scorer.py", "Kết hợp text, entity và constraint", "Thiết kế hợp lý cho retrieval, chưa đủ cho reasoning"),
        ("CACL training", "ours/source/src/gsr_cacl/train.py", "Huấn luyện ba giai đoạn với CHAP negatives", "Đã có khung tốt nhưng chưa nối kín vào benchmark inference"),
        ("Benchmark", "ours/source/src/gsr_cacl/benchmark_gsr.py", "Đánh giá MRR/Recall/NDCG", "Phù hợp cho retrieval, chưa phải benchmark toàn trình"),
    ]
    for row in module_rows:
        cells = table2.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("4. Phân tích kỹ thuật retrieval hiện tại dựa trên contribution1.pdf và mã nguồn", level=1)
    add_paragraphs(
        doc,
        [
            "Phần retrieval trong contribution1.pdf có giá trị đặc biệt vì nó không trình bày GSR như một graph dùng để trang trí, mà như một nỗ lực đưa cấu trúc và ràng buộc kế toán vào đúng pha truy xuất. Tinh thần cốt lõi của proposal này là dense retrieval thuần text thường thất bại trên tài liệu tài chính vì hai lý do: một là hiện tượng nhầm lẫn thực thể khi nhiều báo cáo của cùng công ty qua các năm có bề mặt ngôn ngữ rất giống nhau; hai là hiện tượng làm phẳng bảng khiến mô hình mất thông tin về cấu trúc số học bên trong.",
            "Theo kiến trúc được trình bày trong contribution1.pdf và đối chiếu với code hiện tại, GSR vận hành theo quy trình sau. Câu hỏi và metadata liên quan được dùng để truy xuất một tập ứng viên từ FAISS. Với từng tài liệu ứng viên, hệ trích bảng markdown, dùng template tài chính để chuẩn hóa tiêu đề và xây một Constraint KG, sau đó mã hóa KG này bằng Edge-aware GAT để thu được graph embedding. Song song, hệ tính một constraint score phản ánh mức độ phù hợp của cấu trúc bảng. Cuối cùng, joint scorer kết hợp tín hiệu text, entity và constraint để sắp xếp lại các ứng viên.",
            "Điều cần nhấn mạnh là retrieval ở đây vẫn là retrieval hai tầng. Tầng thứ nhất là dense retrieval để gom candidate; tầng thứ hai mới là reranking có cấu trúc. Nói cách khác, KG hiện chưa thay thế bộ truy xuất chính, mà đóng vai trò là một mô-đun tái chấm điểm ứng viên theo logic tài chính. Đây là một lựa chọn hợp lý cho giai đoạn đầu vì nó cho phép chèn tri thức miền mà không phải tái thiết kế hoàn toàn toàn bộ hệ truy xuất.",
        ],
    )

    doc.add_paragraph()
    doc.add_picture(str(fig_path), width=Inches(6.3))
    cap = doc.add_paragraph("Hình 1. Kiến trúc mô hình: bên trên là pipeline retrieval hiện tại theo tinh thần GSR-CACL; bên dưới là kiến trúc toàn trình được đề xuất cho giai đoạn tiếp theo.", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.1. Constraint KG: từ bảng tài chính đến cấu trúc có hướng", level=2)
    add_paragraphs(
        doc,
        [
            "Constraint KG hiện tại được xây từ bảng markdown bằng ba bước chính. Bước thứ nhất là parse bảng thành hàng và cột, đồng thời chuyển các chuỗi biểu diễn số thành giá trị chuẩn hóa, bao gồm các trường hợp có dấu ngoặc âm, đơn vị triệu, tỷ hoặc phần trăm. Bước thứ hai là chuẩn hóa header tài chính, ví dụ Sales được ánh xạ về Revenue, Net Profit được ánh xạ về Net Income, Tax Expense được ánh xạ về Income Tax. Bước thứ ba là match bảng với một thư viện template IFRS/GAAP để sinh các cạnh accounting hoặc fallback sang các cạnh positional nếu match yếu.",
            "Ý tưởng này có giá trị ở chỗ nó biến bảng từ dạng văn bản phẳng thành một cấu trúc có hướng, nơi mỗi ô là một node và mỗi quan hệ tài chính hoặc quan hệ vị trí là một cạnh. Với cách làm này, mô hình có thể phân biệt tốt hơn giữa những bảng tuy giống nhau về mặt từ vựng nhưng khác nhau về cấu trúc kế toán. Đây chính là lợi thế quan trọng của GSR so với dense retrieval chỉ nhìn vào chuỗi text.",
            "Tuy nhiên, qua phân tích kỹ thuật có thể thấy Constraint KG hiện tại vẫn chủ yếu phù hợp với retrieval hơn là reasoning. Nhiều phương trình kế toán nhiều toán hạng đang bị phân rã thành các cạnh cặp đôi đi vào cùng một node đích. Biểu diễn này tạo được tín hiệu cấu trúc, nhưng chưa đủ trung thành để kiểm chứng một phương trình hay để làm nền cho executor suy luận số học nhiều bước.",
        ],
    )

    doc.add_heading("4.2. Edge-aware GAT: cơ chế tạo graph embedding", level=2)
    add_paragraphs(
        doc,
        [
            "Edge-aware GAT trong hệ hiện tại là mô-đun học được quan trọng nhất ở phía KG. Mỗi node trước hết được mã hóa từ ba loại thông tin: nhãn header đã chuẩn hóa, đặc trưng số của giá trị trong ô và mã hóa vị trí theo hàng và cột. Từ đó, mạng GAT lan truyền thông tin dọc theo các cạnh của đồ thị để tạo ra embedding cho từng node và sau cùng gộp chúng thành một graph embedding cấp tài liệu.",
            "Điểm khác biệt quan trọng so với GAT thông thường là trọng số cạnh omega không chỉ tồn tại như một nhãn phụ mà đi trực tiếp vào cơ chế attention. Trong triển khai hiện tại, omega vừa tạo bias cho attention, vừa ảnh hưởng đến message passing giữa node nguồn và node đích. Điều này làm cho cạnh cộng, cạnh trừ và cạnh vị trí mang ý nghĩa khác nhau trong quá trình mã hóa. Nhờ vậy, KG không chỉ nói rằng hai ô có liên quan, mà còn nói chúng liên quan theo loại phép toán nào.",
            "Ở góc nhìn retrieval, điều này rất hữu ích vì graph embedding sẽ mang thêm dấu vết của cấu trúc bảng thay vì chỉ mang nghĩa từ vựng. Trong các tình huống có nhiều báo cáo gần giống nhau, sự khác biệt nhỏ về quan hệ hàng, cột, tổng và thành phần có thể là tín hiệu đủ mạnh để tách đúng tài liệu khỏi các distractor khó. Tuy nhiên, hiệu quả này vẫn phụ thuộc mạnh vào việc template matching có tốt hay không, bởi nếu đồ thị rơi vào fallback positional edges thì tín hiệu học được sẽ yếu hơn đáng kể.",
        ],
    )

    doc.add_heading("4.3. Constraint scoring: tín hiệu kiểm tra cấu trúc", level=2)
    add_paragraphs(
        doc,
        [
            "Bên cạnh graph embedding học được, hệ còn xây một tín hiệu thủ công gọi là constraint score. Về trực giác, đây là một chỉ báo cho biết một bảng ứng viên có vẻ nhất quán về mặt cấu trúc tài chính hay không. Tín hiệu này được tính bằng cách duyệt các cạnh accounting trong đồ thị, đo sai lệch giữa node nguồn và node đích theo omega, sau đó chuyển sai lệch đó qua một hàm mũ để thu về điểm trong khoảng từ 0 đến 1.",
            "Ưu điểm của constraint score là nó bổ sung một prior tài chính tương đối trực tiếp vào quá trình reranking. Trong nhiều trường hợp, hai tài liệu có thể rất giống nhau về semantic similarity, nhưng tài liệu đúng sẽ có cấu trúc số liệu hợp lý hơn. Constraint score giúp mô hình tận dụng được khác biệt này. Đây là một trong những điểm mạnh thực dụng nhất của kiến trúc retrieval hiện tại.",
            "Dù vậy, constraint score cũng là nơi bộc lộ rõ nhất giới hạn của thiết kế hiện tại. Vì các ràng buộc nhiều toán hạng đang được quy về các cạnh cặp đôi, điểm số sinh ra chỉ là một phép xấp xỉ, không phải kiểm chứng phương trình đúng nghĩa. Hệ quả là bảng đúng chưa chắc đạt điểm tối đa, còn bảng không match template mạnh đôi khi lại không bị phạt đủ rõ. Điều này củng cố nhận định rằng constraint score hiện tại có giá trị trong retrieval, nhưng không thể là bộ kiểm định chính cho reasoning.",
        ],
    )

    doc.add_heading("4.4. Joint scorer và vai trò thực của metadata", level=2)
    add_paragraphs(
        doc,
        [
            "Joint scorer là nơi ba nguồn tín hiệu được hợp nhất: text similarity, entity score và constraint signal. Đây là phần rất quan trọng vì nó cho thấy GSR không phủ nhận semantic retrieval, mà chỉ đặt semantic retrieval vào một khung scoring rộng hơn có tri thức miền. Nhờ cách hợp nhất này, hệ có thể ưu tiên các tài liệu vừa giống câu hỏi về nội dung, vừa đúng công ty, đúng năm và có cấu trúc tài chính hợp lý.",
            "Metadata hiện tại gồm company, year và sector. Ý tưởng ban đầu là dùng chúng như tín hiệu phân biệt thực thể, đặc biệt để khắc phục hiện tượng cùng công ty nhưng khác năm. Đây là một ý tưởng đúng, nhưng qua phân tích có thể thấy metadata vẫn chưa phát huy hết vai trò. Trong hệ hiện tại, metadata mới chủ yếu tham gia ở mức scoring bề mặt, chứ chưa thật sự đi vào biểu diễn chunk, vào hard-negative mining theo kiểu ontology, hay vào reasoning như một điều kiện tương thích của toán hạng.",
            "Điểm này đặc biệt quan trọng cho giai đoạn tiếp theo. Nếu coi metadata chỉ là bonus score, đóng góp của nó sẽ bị giới hạn. Nếu coi metadata là một ontology nhẹ của bài toán tài chính, nó có thể tham gia vào toàn bộ pipeline: điều kiện định tuyến candidate, cấu trúc chunk embedding, điều kiện triage evidence, ràng buộc chọn toán hạng và cả verifier hậu nghiệm. Đây là một hướng mở rất giàu tiềm năng mà hệ hiện tại mới chạm đến ở mức sơ khai.",
        ],
    )

    doc.add_heading("5. CACL và ý nghĩa của học tăng cường / tối ưu ưu tiên", level=1)
    add_paragraphs(
        doc,
        [
            "Trong contribution1.pdf, CACL được đưa ra như cơ chế huấn luyện nhằm làm cho retriever và scorer phân biệt được các tài liệu gần đúng nhưng sai về logic tài chính. Điểm hay của ý tưởng này nằm ở CHAP negatives: thay vì chỉ chọn negative khác chủ đề, hệ chủ động tạo các negative khó bằng cách làm sai một công thức, đổi scale hoặc đổi metadata thực thể. Đây là một triết lý huấn luyện rất đáng giữ lại vì nó phù hợp với đúng loại lỗi mà hệ retrieval tài chính thường mắc phải.",
            "Tuy nhiên, nếu đi sang bài toán reasoning toàn trình thì CACL theo dạng hiện tại vẫn chưa đủ. Lý do là lỗi của reasoning không còn chỉ nằm ở việc chọn sai tài liệu, mà nằm ở việc chọn sai toán hạng trong tập context đã truy xuất, chuẩn hóa sai đơn vị hoặc sinh sai chương trình tính toán. Khi đó, huấn luyện chỉ dựa vào contrastive retrieval sẽ không tác động trực tiếp vào phần quyết định nhất của bài toán.",
            "Vì vậy, chiến lược hợp lý hơn là xem CACL như nền retrieval-oriented, sau đó bổ sung một tầng huấn luyện mới cho reasoning. Tầng này nên kết hợp Supervised Fine-Tuning cho chương trình suy luận, preference optimization ở mức trace và reinforcement learning với reward có thể kiểm chứng bằng executor. Đây là lý do phần sau của báo cáo đề xuất pipeline SFT rồi Step-DPO rồi GRPO thay vì chỉ mở rộng thêm contrastive learning thuần retrieval.",
        ],
    )

    doc.add_heading("5.1. So sánh DPO, ORPO và GRPO cho bài toán hiện tại", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Để lựa chọn thuật toán huấn luyện phù hợp, cần đặt câu hỏi theo đặc thù bài toán tài chính chứ không theo độ phổ biến chung của từng thuật toán. Bảng dưới đây tổng hợp vai trò phù hợp của các lựa chọn chính.")

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = "Table Grid"
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers3 = ["Thuật toán", "Ưu điểm", "Hạn chế", "Vai trò đề xuất trong đề tài"]
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "F4CCCC")
    set_repeat_table_header(table3.rows[0])

    algo_rows = [
        ("DPO", "Ổn định, offline, dễ triển khai hơn RL online", "Tối ưu ở mức toàn câu trả lời, chưa khai thác mạnh reward thực thi", "Dùng sau SFT để chỉnh reasoning trace và tránh context nhiễu"),
        ("ORPO", "Chi phí thấp, phù hợp làm baseline", "Khó trở thành tối ưu chính cho numerical reasoning có executor", "Dùng như baseline so sánh compute thấp"),
        ("Step-DPO", "Nhấn mạnh chất lượng từng bước suy luận", "Cần chuẩn bị cặp trace chất lượng", "Lựa chọn ưu tiên hơn DPO thuần cho trace reasoning"),
        ("GRPO / RLVR", "Phù hợp bài toán có reward tính tự động từ executor", "Phụ thuộc mạnh vào chất lượng reward và verifier", "Thuật toán chính cho tối ưu reasoning cuối cùng"),
    ]
    for row in algo_rows:
        cells = table3.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    add_paragraphs(
        doc,
        [
            "Sau khi đối chiếu các thuật toán với bản chất bài toán tài chính, kết luận hợp lý nhất là không nên dùng ORPO như chiến lược chính và cũng không nên dừng ở DPO thuần. DPO có ích ở pha trung gian để dạy mô hình phân biệt trace grounded với trace dùng context nhiễu, nhưng nếu bài toán đã có executor và reward có thể kiểm chứng thì GRPO hoặc một biến thể RLVR mới là lựa chọn mạnh cho giai đoạn tối ưu cuối. Đồng thời, do lỗi reasoning thường nằm ở một bước trung gian như sai year hoặc sai unit, Step-DPO hoặc một biến thể process-aware sẽ phù hợp hơn response-level DPO.",
        ],
    )

    doc.add_heading("6. Chốt ý tưởng nghiên cứu và kiến trúc đề xuất", level=1)
    add_paragraphs(
        doc,
        [
            "Sau khi phân tích benchmark, code hiện tại và tài liệu liên quan, ý tưởng nên chốt không phải là một KG tốt hơn cho retrieval, mà là một Financial Evidence Graph dùng chung cho retrieval và reasoning. Đồ thị này cần liên kết được document, section, table, row, column, cell, sentence, footnote, company, period, metric concept, unit và equation. Chỉ khi đó KG mới thực sự trở thành cầu nối giữa truy xuất và suy luận, thay vì dừng ở vai trò structural reranker.",
            "Trong kiến trúc đề xuất, retrieval cần chuyển sang dạng phân cấp. Thay vì chỉ lấy top-k document rồi đưa nguyên văn cho mô hình suy luận, hệ sẽ đi qua ba tầng: document retrieval, table/section retrieval và evidence atom retrieval. Evidence atom ở đây có thể là cell, row aggregate, sentence hoặc footnote span. Từ các atom này, hệ mới xây local evidence graph cho từng câu hỏi cụ thể. Đây là khác biệt quyết định giữa một hệ retrieval-centric và một hệ evidence-grounded reasoning.",
            "Reasoning sau đó không chạy trên raw context mà chạy trên local evidence graph đã được triage. Câu hỏi sẽ được parser để xác định company, phạm vi thời gian, metric mục tiêu và loại phép toán. Module grounding sẽ chọn các toán hạng tương thích, sau đó một planner sẽ sinh DSL hoặc chương trình Python, executor thực thi phép tính và verifier kiểm tra ngược xem đáp án có nhất quán với company, year, đơn vị, scale và equation constraints hay không. Cách tiếp cận này biến bài toán trả lời câu hỏi thành bài toán dựng và kiểm chứng cấu trúc bằng chứng, phù hợp hơn với miền tài chính.",
        ],
    )

    doc.add_heading("6.1. Vì sao hướng này mạnh hơn việc chỉ cải thiện MRR", level=2)
    add_paragraphs(
        doc,
        [
            "Trong tài liệu tài chính, tăng MRR là cần nhưng chưa đủ. Một hệ có MRR@3 tốt nhưng thường xuyên trộn một số từ context đúng với một số từ context nhiễu vẫn không giải được bài toán thực. Điều advisor hoặc hội đồng đánh giá thực sự quan tâm là liệu hệ có đáng tin hay không khi đối diện tình huống retrieval không hoàn hảo. Kiến trúc dựa trên evidence graph và verifier trả lời đúng vào điểm đó, bởi nó không giả định rằng top-k context đã sạch mà chủ động giải quyết nhiễu trước và trong reasoning.",
            "Hướng mới cũng mạnh hơn về mặt câu chuyện học thuật. Nếu bài báo chỉ nói KG giúp retrieval tốt hơn thì rất dễ bị nhìn như một biến thể của graph reranking. Nhưng nếu bài báo cho thấy cùng một evidence graph được dùng để chọn bằng chứng, neo toán hạng, sinh chương trình và kiểm chứng kết quả thì contribution trở nên sâu hơn nhiều. Nó không còn là một mẹo retrieval mà là một cơ chế hợp nhất retrieval và reasoning trong miền tài chính.",
        ],
    )

    doc.add_heading("7. Kế hoạch triển khai chi tiết", level=1)
    add_paragraphs(
        doc,
        [
            "Lộ trình triển khai nên được chia thành các chặng rõ ràng để bảo đảm mỗi bước đều tạo ra kết quả trung gian có thể đánh giá và báo cáo. Việc này đặc biệt quan trọng vì bài toán toàn trình có độ rộng lớn; nếu không chia chặng hợp lý, hệ thống rất dễ rơi vào trạng thái nhiều ý tưởng nhưng không có mô-đun nào đủ chín.",
        ],
    )

    table4 = doc.add_table(rows=1, cols=4)
    table4.style = "Table Grid"
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers4 = ["Giai đoạn", "Mục tiêu", "Công việc chính", "Sản phẩm mong đợi"]
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9D2E9")
    set_repeat_table_header(table4.rows[0])

    roadmap_rows = [
        ("Giai đoạn 1", "Củng cố retrieval", "Mở rộng metadata schema, contextual chunk embedding, hard negatives mạnh hơn, table/section reranking", "Retrieval phân cấp mạnh hơn và có ablation rõ"),
        ("Giai đoạn 2", "Xây evidence graph", "Thiết kế schema node/edge, provenance, unit/scale, equation node, liên kết text-table-footnote", "Financial Evidence Graph hoạt động được"),
        ("Giai đoạn 3", "Xây reasoning substrate", "Query parser, operand grounding, DSL/Python executor, verifier", "Pipeline reasoning có thể thực thi và kiểm chứng"),
        ("Giai đoạn 4", "Huấn luyện reasoning", "SFT, Step-DPO, GRPO/RLVR với reward verifiable", "Mô hình reasoning ổn định và đáng tin hơn"),
        ("Giai đoạn 5", "Đánh giá toàn trình", "End-to-end benchmark, robustness dưới top-3 noisy contexts, case study", "Bộ kết quả hoàn chỉnh cho báo cáo và bài báo"),
    ]
    for row in roadmap_rows:
        cells = table4.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    add_paragraphs(
        doc,
        [
            "Thứ tự triển khai trong bảng trên không chỉ mang tính quản lý dự án mà còn phản ánh phụ thuộc kỹ thuật giữa các mô-đun. Retrieval phải đủ tốt thì evidence triage mới có đầu vào có ích; evidence graph phải rõ thì grounding và verifier mới đứng được; verifier phải đáng tin thì reward của RL mới không bị méo. Chính vì vậy, đề xuất huấn luyện SFT rồi Step-DPO rồi GRPO chỉ có ý nghĩa sau khi executor và verifier đã ổn định.",
        ],
    )

    doc.add_heading("8. Đóng góp khoa học dự kiến", level=1)
    add_paragraphs(
        doc,
        [
            "Nếu triển khai thành công, công trình có thể chốt ở ba đóng góp khoa học rõ ràng. Đóng góp thứ nhất là metadata-aware hierarchical retrieval cho tài liệu tài chính dài, trong đó metadata được nâng thành ontology truy xuất và retrieval được đẩy xuống cấp table/section rồi evidence atom. Đóng góp thứ hai là Financial Evidence Graph trung thành với phương trình, dùng chung cho retrieval và reasoning thay vì chỉ phục vụ structural reranking. Đóng góp thứ ba là pipeline reasoning có thể kiểm chứng, kết hợp grounding, executor, verifier và tối ưu theo hướng Step-DPO rồi GRPO.",
            "Ba đóng góp này có mối liên hệ tự nhiên với nhau. Retrieval phân cấp giải quyết việc đưa đúng bằng chứng vào hệ; evidence graph tạo không gian biểu diễn chung cho cả retrieval và reasoning; còn executor cùng verifier giúp chuyển từ RAG trả lời ngôn ngữ sang reasoning số học đáng tin. Chính sự gắn kết này làm cho đề tài có chiều sâu hơn một bài retrieval thuần hoặc một bài reasoning tách biệt khỏi truy xuất.",
        ],
    )

    doc.add_heading("9. Kết luận", level=1)
    add_paragraphs(
        doc,
        [
            "Báo cáo này đi từ hiện trạng triển khai thực tế trong repo, nội dung retrieval trong contribution1.pdf, các benchmark đang dùng, những khoảng trống kỹ thuật còn tồn tại, cho đến một chiến lược triển khai toàn trình đủ rõ để có thể báo cáo và làm việc tiếp theo với giáo viên hướng dẫn. Kết luận tổng quát là đề tài hiện đã có một prototype retrieval có cấu trúc đủ tốt để làm nền, nhưng chưa nên dừng ở đó.",
            "Hướng phát triển mạnh nhất là biến KG hiện tại thành Financial Evidence Graph dùng chung cho retrieval và reasoning, đồng thời nâng metadata thành ontology truy xuất và thêm một chuỗi reasoning có executor cùng verifier. Trên nền tảng đó, chiến lược huấn luyện phù hợp nhất là retrieval ranking ở giai đoạn đầu, SFT cho reasoning, Step-DPO để làm sạch trace và GRPO để tối ưu cuối cùng bằng reward verifiable. Đây là lộ trình vừa khả thi về mặt kỹ thuật, vừa có câu chuyện học thuật đủ rõ để phát triển thành kết quả nghiên cứu mạnh.",
        ],
    )

    doc.add_heading("Tài liệu tham khảo chính", level=1)
    refs = [
        "Rafailov et al., Direct Preference Optimization: Your Language Model is Secretly a Reward Model, 2023. https://arxiv.org/abs/2305.18290",
        "Hong et al., ORPO: Monolithic Preference Optimization without Reference Model, EMNLP 2024. https://aclanthology.org/2024.emnlp-main.626/",
        "Shao et al., DeepSeekMath: Pushing the Limits of Mathematical Reasoning in Open Language Models, 2024. https://arxiv.org/abs/2402.03300",
        "Chen et al., FinQA: A Dataset of Numerical Reasoning over Financial Data, EMNLP 2021. https://aclanthology.org/2021.emnlp-main.300/",
        "Zhu et al., TAT-QA: A Question Answering Benchmark on a Hybrid of Tabular and Textual Content in Finance, ACL 2021. https://aclanthology.org/2021.acl-long.254/",
        "Chen et al., ConvFinQA: Exploring the Chain of Numerical Reasoning in Conversational Finance Question Answering, EMNLP 2022. https://aclanthology.org/2022.emnlp-main.421/",
        "T2-RAGBench: Retrieval-first Benchmark for Text-and-Table Financial QA, 2025 preprint. https://arxiv.org/abs/2506.12071",
        "HierFinRAG: Hierarchical Financial Retrieval-Augmented Generation, 2026. https://www.mdpi.com/2227-9709/13/2/30",
        "FT-RAG: Fine-Grained Table Retrieval-Augmented Generation for Financial Analysis, 2026 preprint. https://arxiv.org/abs/2605.01495",
        "Metadata-Driven Financial RAG: Retrieval-Enhanced Large Language Models for Financial Question Answering, 2025 preprint. https://arxiv.org/abs/2510.24402",
        "Structure First, Reason Next: A Framework for Financial Numerical Reasoning with Structured Evidence, 2026 preprint. https://arxiv.org/abs/2601.07754",
        "Table-R1: Reinforcement Learning for Table Reasoning, 2025 preprint. https://arxiv.org/abs/2505.23621",
        "Contribution nội bộ: Structured Knowledge-Enhanced Retrieval for Financial Documents, contribution1.pdf trong thư mục NLP.",
    ]
    add_reference_list(doc, refs)

    doc.save(docx_path)


def build_progress_report(docx_path: Path, fig_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    configure_styles(doc)
    add_header_footer(doc)

    core = doc.core_properties
    core.title = "Báo cáo tiến độ nghiên cứu KG-assisted financial reasoning"
    core.author = "OpenAI Codex"
    core.subject = "Báo cáo tiến độ nghiên cứu gửi GVHD"

    fig1 = CONTRIB_DIR / "figure1_pipeline_crop.png"
    fig2 = CONTRIB_DIR / "figure2_kg_construction.png"
    fig3 = CONTRIB_DIR / "figure3_gat_overview.png"
    fig4 = CONTRIB_DIR / "figure4_edge_aware_attention.png"
    fig5 = CONTRIB_DIR / "figure5_gat_encoder_flow.png"
    fig6 = CONTRIB_DIR / "figure6_joint_scorer_crop.png"
    fig7 = CONTRIB_DIR / "figure7_chap_sampler.png"
    fig8 = CONTRIB_DIR / "figure8_cacl_objective.png"
    fig9 = CONTRIB_DIR / "figure9_curriculum.png"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(78)
    r = p.add_run("BÁO CÁO TIẾN ĐỘ NGHIÊN CỨU\nRETRIEVAL VÀ SUY LUẬN TÀI CHÍNH DỰA TRÊN ĐỒ THỊ TRI THỨC")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor(18, 60, 110)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(18)
    sr = sub.add_run(
        "Tổng hợp hiện trạng triển khai GSR, CACL và Constraint KG; "
        "mô tả benchmark T²-RAGBench; đánh giá khách quan các khoảng trống kỹ thuật; "
        "và chốt định hướng Financial Evidence Graph dùng chung cho retrieval và reasoning"
    )
    sr.font.name = "Times New Roman"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    sr.font.size = Pt(12.5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(122)
    mr = meta.add_run("Ngày cập nhật: 13/06/2026\nTài liệu phục vụ báo cáo tiến độ với giáo viên hướng dẫn")
    mr.font.name = "Times New Roman"
    mr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    mr.font.size = Pt(12)

    doc.add_page_break()

    doc.add_heading("1. Tóm tắt điều hành", level=1)
    add_paragraphs(
        doc,
        [
            "Mục tiêu của báo cáo này là chuyển phần ý tưởng và phần triển khai hiện có thành một bức tranh tiến độ nhất quán, trong đó làm rõ ba lớp nội dung: một là bài toán nghiên cứu thực sự cần giải; hai là những gì đã được hiện thực hóa trong repo hiện tại; ba là định hướng kỹ thuật đủ mạnh để nối retrieval với reasoning thành một pipeline tài chính toàn trình.",
            "Qua đối chiếu trực tiếp giữa contribution1.pdf, mã nguồn trong thư mục NLP/ours/source và báo cáo tiến độ dạng markdown hiện có, có thể kết luận rằng đề tài đã sở hữu một prototype retrieval có cấu trúc khá rõ. Ba hạt nhân kỹ thuật đã xuất hiện tương đối đầy đủ là Graph-Structured Retrieval (GSR), Constraint-Aware Contrastive Learning (CACL) và Constraint KG xây từ bảng markdown. Tuy nhiên, hệ thống hiện nay vẫn thiên về document retrieval và reranking, còn reasoning số học đa bước sau retrieval vẫn chưa được khép kín.",
            "Điểm quan trọng nhất cần chốt ở giai đoạn này là: đồ thị tri thức hiện tại khả thi và hữu ích cho retrieval, nhưng chưa đủ để trực tiếp gánh reasoning số học toàn trình. Vì vậy, định hướng phù hợp không phải chỉ tăng thêm tín hiệu graph cho reranking, mà là nâng Constraint KG thành một Financial Evidence Graph có kiểu, có provenance, có ontology metadata và có khả năng phục vụ đồng thời việc truy xuất, neo toán hạng, thực thi chương trình và kiểm chứng kết quả.",
        ],
    )

    doc.add_heading("2. Mô tả bài toán và benchmark T²-RAGBench", level=1)
    add_paragraphs(
        doc,
        [
            "Bài toán đặt ra không còn là truy xuất ngữ cảnh theo nghĩa RAG thông thường, mà là giải quyết toàn trình một câu hỏi tài chính trong điều kiện dữ liệu đầu vào gồm báo cáo dài, nhiều bảng, nhiều đoạn văn và nhiều footnote. Trong setting này, hệ thống phải tự tìm bằng chứng liên quan, tự chọn đúng toán hạng, tự thực hiện phép tính và có cơ chế tự kiểm tra tính nhất quán của kết quả cuối cùng.",
            "T²-RAGBench là benchmark phù hợp với mục tiêu đó vì nó buộc hệ thống phải retrieval trước khi reasoning, thay vì cung cấp sẵn oracle context. Theo công bố gốc trên arXiv, benchmark này gồm 32.908 bộ ba câu hỏi - ngữ cảnh - đáp án trên tài liệu tài chính pha trộn văn bản và bảng, đồng thời được thiết kế để tránh hiện tượng một câu hỏi có nhiều đáp án đúng tùy theo context được cấp. Điểm này rất quan trọng vì nó làm benchmark trở nên phù hợp cho đánh giá retrieval-augmented reasoning trong bối cảnh thực tế hơn.",
        ],
    )

    io_table = doc.add_table(rows=1, cols=2)
    io_table.style = "Table Grid"
    io_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(io_table.rows[0].cells[0], "Thành phần", bold=True, font_size=10)
    set_cell_text(io_table.rows[0].cells[1], "Mô tả trong bài toán toàn trình", bold=True, font_size=10)
    shade_cell(io_table.rows[0].cells[0], "D9EAF7")
    shade_cell(io_table.rows[0].cells[1], "D9EAF7")
    set_repeat_table_header(io_table.rows[0])
    io_rows = [
        ("Đầu vào", "Một truy vấn Q và một corpus tài chính C = {Ci}. Mỗi tài liệu Ci chứa văn bản tường thuật ti, bảng markdown Ti và metadata mi như company, year, sector."),
        ("Đầu ra retrieval", "Danh sách top-k bằng chứng hoặc ứng viên trung gian, không chỉ ở cấp tài liệu mà lý tưởng là ở cấp section, table và evidence atom."),
        ("Đầu ra end-to-end", "Đáp án số cuối cùng y, kèm theo bằng chứng được grounding, dấu vết suy luận hoặc chương trình thực thi, và trạng thái kiểm chứng hợp lệ/không hợp lệ."),
        ("Yêu cầu toàn trình", "Hệ thống phải đúng về công ty, năm, đơn vị, scale, đúng toán hạng và đúng phép toán; đồng thời chịu được trường hợp top-k còn nhiễu."),
        ("Đánh giá cần có", "MRR/Recall/NDCG cho retrieval; độ chính xác evidence grounding; answer accuracy và execution accuracy cho reasoning; robustness dưới context nhiễu."),
    ]
    for left, right in io_rows:
        row = io_table.add_row().cells
        set_cell_text(row[0], left, font_size=10)
        set_cell_text(row[1], right, font_size=10)

    doc.add_heading("2.1. Ý nghĩa của T²-RAGBench đối với đề tài", level=2)
    add_paragraphs(
        doc,
        [
            "Nếu chỉ dùng các benchmark như FinQA ở chế độ đã có sẵn ngữ cảnh đúng, mô hình rất dễ cho cảm giác reasoning tốt trong khi thực tế lại thất bại ngay từ bước retrieval. T²-RAGBench ép hệ thống phải đối mặt với hiện tượng đặc trưng của miền tài chính: các báo cáo của cùng một công ty ở nhiều năm có ngôn ngữ rất giống nhau; nhiều bảng có cấu trúc gần nhau; và câu trả lời cuối cùng thường phụ thuộc vào đúng một vài ô số trong rất nhiều distractor nhìn bề ngoài khá hợp lý.",
            "Vì vậy, benchmark này giúp nhìn rõ một sự thật quan trọng: tăng MRR@3 là cần nhưng chưa đủ. Một hệ thống có thể đưa tài liệu đúng vào top-3 nhưng vẫn thất bại nếu reasoning module chọn nhầm hai toán hạng nhiễu còn lại. Đây cũng là lý do vì sao hướng phát triển tiếp theo phải đi từ document retrieval sang evidence grounding và operand grounding.",
        ],
    )

    ds_table = doc.add_table(rows=1, cols=4)
    ds_table.style = "Table Grid"
    ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ds_headers = ["Bộ dữ liệu", "Vai trò trong nghiên cứu", "Đặc trưng nổi bật", "Cách dùng đề xuất"]
    for i, header_text in enumerate(ds_headers):
        set_cell_text(ds_table.rows[0].cells[i], header_text, bold=True, font_size=10)
        shade_cell(ds_table.rows[0].cells[i], "E2F0D9")
    set_repeat_table_header(ds_table.rows[0])
    for row_data in [
        ("T²-RAGBench", "Benchmark retrieval-first", "Không giả định oracle context; phù hợp cho text và table", "Benchmark chính cho retrieval và đánh giá end-to-end sau này"),
        ("FinQA", "Chuẩn reasoning số học", "Có gold program và bài toán nhiều bước", "Dùng huấn luyện executor/planner và đánh giá tính đúng chương trình"),
        ("ConvFinQA", "Mở rộng reasoning", "Giữ ngữ cảnh hội thoại và chain dài", "Dùng stress-test reasoning dưới truy vấn nhiều lượt"),
        ("TAT-QA / DocFinQA", "Bài toán hybrid table-text", "Cần grounding đồng thời bảng và văn bản", "Dùng đánh giá grounding và robustness đa nguồn"),
    ]:
        cells = ds_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("3. Hiện trạng triển khai trong repo", level=1)
    add_paragraphs(
        doc,
        [
            "Qua khảo sát mã nguồn và tài liệu trong thư mục NLP/, có thể xác định rằng đề tài đã đi qua giai đoạn ý tưởng sơ khai và đang ở mức một prototype nghiên cứu thực sự. Repo hiện có cả nhánh baseline retrieval và nhánh đề xuất, trong đó nhánh đề xuất đã triển khai tương đối đầy đủ các thành phần chính của GSR-CACL.",
            "Điểm mạnh của giai đoạn hiện tại là các thành phần không nằm rời rạc ở mức ghi chú, mà đã xuất hiện thành code chạy được: có dataset wrapper cho benchmark, có builder sinh KG từ bảng markdown, có GAT encoder, có joint scorer, có negative sampler kiểu CHAP và có benchmark script để đo MRR/Recall/NDCG. Điều này cho thấy luận điểm trong contribution1.pdf đã có nền tảng thực thi chứ không chỉ là đề cương ý tưởng.",
        ],
    )

    repo_table = doc.add_table(rows=1, cols=4)
    repo_table.style = "Table Grid"
    repo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repo_headers = ["Thành phần", "Đường dẫn chính", "Đã triển khai gì", "Nhận xét tiến độ"]
    for i, header_text in enumerate(repo_headers):
        set_cell_text(repo_table.rows[0].cells[i], header_text, bold=True, font_size=10)
        shade_cell(repo_table.rows[0].cells[i], "FCE5CD")
    set_repeat_table_header(repo_table.rows[0])
    repo_rows = [
        ("Dataset wrapper", "ours/source/src/gsr_cacl/datasets/wrappers.py", "Nạp T²-RAGBench, tổ chức corpus và metadata", "Đã dùng được; là nền cho benchmark hiện tại"),
        ("GSR retrieval", "ours/source/src/gsr_cacl/methods/gsr_retrieval.py", "Lấy ứng viên bằng FAISS và rerank bằng scorer có cấu trúc", "Đây là phần inference rõ nhất của hệ hiện tại"),
        ("KG builder", "ours/source/src/gsr_cacl/kg/builder.py", "Tạo Constraint KG từ bảng markdown bằng template và fallback positional edges", "Khả thi cho retrieval; chưa đủ trung thành cho reasoning"),
        ("GAT layer / encoder", "ours/source/src/gsr_cacl/encoders/gat_layer.py", "Lan truyền thông tin có nhận biết omega trên cạnh", "Đã hiện thực hóa đúng tinh thần edge-aware GAT"),
        ("Constraint scoring", "ours/source/src/gsr_cacl/scoring/constraint_score.py", "Tính điểm tuân thủ ràng buộc từ residual trên accounting edges", "Hữu ích cho reranking nhưng còn xấp xỉ"),
        ("Joint scorer", "ours/source/src/gsr_cacl/scoring/joint_scorer.py", "Kết hợp s_text, s_entity, s_constraint với trọng số học được", "Đã có cả mode train và score_single cho inference"),
        ("CHAP sampler", "ours/source/src/gsr_cacl/negative_sampler/chap.py", "Sinh negative kiểu additive, scale, entity/year swap", "Đây là đóng góp training quan trọng của CACL"),
        ("Training pipeline", "ours/source/src/gsr_cacl/train.py", "Ba giai đoạn Identity, Structural, Joint CACL", "Có khung train; chưa nối kín với benchmark end-to-end"),
        ("Benchmark", "ours/source/src/gsr_cacl/benchmark_gsr.py", "Đo MRR@3, Recall@k, NDCG@3", "Phù hợp cho retrieval; chưa đo reasoning"),
    ]
    for row_data in repo_rows:
        cells = repo_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("4. Retrieval hiện tại dựa trên contribution1.pdf và mã nguồn", level=1)
    add_paragraphs(
        doc,
        [
            "Phần retrieval trong contribution1.pdf là đóng góp đã hiện hình rõ nhất của đề tài. Tinh thần cốt lõi của proposal này là dense retrieval thuần text chưa đủ cho tài liệu tài chính, vì nó bỏ qua cả metadata thực thể lẫn cấu trúc số học nội tại của bảng. GSR khắc phục điểm đó bằng cách đưa thêm hai lớp tín hiệu vào giai đoạn reranking: tín hiệu graph và tín hiệu constraint.",
            "Ở mức luồng xử lý, query và tài liệu trước hết được mã hóa để lấy tín hiệu ngôn ngữ. Song song, metadata được tách thành một tín hiệu entity riêng. Với mỗi ứng viên tài liệu, hệ thống trích bảng markdown, dựng Constraint KG, chạy Edge-aware GAT để thu graph embedding và tính Constraint Score để phản ánh mức nhất quán kế toán. Ba tín hiệu sau đó được hợp nhất trong Joint Scorer để tái xếp hạng ứng viên.",
        ],
    )
    add_captioned_picture(
        doc,
        fig1,
        "Hình 1. Figure 1 trong contribution1.pdf: kiến trúc tổng quát GSR-CACL, cho thấy rõ ba nhánh tín hiệu text, entity và structure cùng đi vào Joint Scorer.",
        Inches(6.1),
    )
    add_display_equation(doc, "s(Q, C) = α · s_text(Q, C) + β · s_ent(Q, C) + γ · CS(G_D)")

    signal_table = doc.add_table(rows=1, cols=3)
    signal_table.style = "Table Grid"
    signal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Tín hiệu", "Vai trò", "Hiện thực trong code"]):
        set_cell_text(signal_table.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(signal_table.rows[0].cells[i], "D9EAD3")
    set_repeat_table_header(signal_table.rows[0])
    signal_rows = [
        ("Text signal", "Đo tương đồng ngữ nghĩa giữa truy vấn và tài liệu", "Text encoder và cosine similarity trong joint_scorer.py"),
        ("Entity signal", "Phân biệt đúng công ty, đúng năm, đúng ngành", "compute_entity_score và forward_entity trong constraint_score.py / joint_scorer.py"),
        ("Structure signal", "Bổ sung tri thức kế toán và tính nhất quán bảng", "KG builder, gat_layer.py và constraint_score.py"),
    ]
    for row_data in signal_rows:
        cells = signal_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("4.1. Constraint KG: biểu diễn bảng tài chính thành đồ thị có hướng", level=2)
    add_paragraphs(
        doc,
        [
            "Constraint KG hiện tại được xây từ bảng markdown qua ba thao tác chính. Trước hết, bảng được parse thành hàng, cột và ô số; tiếp theo các header được chuẩn hóa theo template IFRS/GAAP; cuối cùng hệ thống sinh các accounting edges nếu phát hiện được đẳng thức tài chính, hoặc fallback sang positional edges nếu độ tin cậy match chưa đủ cao.",
            "Vai trò của KG trong giai đoạn này là biến bảng từ một chuỗi markdown phẳng thành một cấu trúc có nghĩa. Khi mỗi ô trở thành node, còn quan hệ cộng, trừ hoặc quan hệ vị trí trở thành cạnh có hướng, mô hình có thêm khả năng phân biệt các bảng bề mặt rất giống nhau nhưng thực chất khác logic kế toán. Đây chính là điểm KG tạo ra giá trị rõ nhất cho retrieval.",
            "Tuy nhiên, cần đánh giá khách quan rằng KG hiện tại chủ yếu là retrieval graph. Nó hữu ích để tái xếp hạng ứng viên tài liệu, nhưng chưa phải là evidence graph giàu ngữ nghĩa đủ để biểu diễn trọn vẹn quan hệ giữa table, text, footnote, đơn vị và provenance. Vì vậy nó là bước đúng, nhưng mới là bước đầu.",
        ],
    )
    add_captioned_picture(
        doc,
        fig2,
        "Hình 2. Figure 2 trong contribution1.pdf: quy trình Parse → Template Matching → Constraint Edge Construction để tạo Constraint KG từ bảng markdown.",
        Inches(5.6),
    )
    add_display_equation(
        doc,
        "conf(H, τ) = |{h ∈ H | normalize(h) ∈ H_τ}| / max(|H|, |H_τ|)"
    )

    doc.add_heading("4.2. Edge-aware GAT: cách KG tạo ra graph embedding", level=2)
    add_paragraphs(
        doc,
        [
            "Sau khi có Constraint KG, hệ thống dùng một GAT layer có nhận biết cạnh để lan truyền thông tin trên đồ thị. Điểm khác của mô-đun này so với GAT thông thường là trọng số cạnh omega không chỉ tồn tại như nhãn bổ trợ, mà đi trực tiếp vào cả attention bias lẫn message passing. Điều đó giúp mô hình phân biệt được cạnh cộng, cạnh trừ và cạnh vị trí.",
            "Trong code hiện tại, gat_layer.py hiện thực hóa đúng tinh thần đó: edge_weight trước hết được chiếu qua edge_proj để sinh edge_bias cho attention; sau đó lại tiếp tục nhân vào message khi node nguồn truyền thông tin sang node đích. Vì vậy, embedding cuối cùng của đồ thị mang theo dấu vết của cấu trúc tài chính, thay vì chỉ là tổng hợp thuần văn bản.",
        ],
    )
    add_captioned_picture(
        doc,
        fig3,
        "Hình 3. Figure 3 trong contribution1.pdf: tổng quan khối GAT Encoder dùng để biến node features thành graph embedding.",
        Inches(5.8),
    )
    add_captioned_picture(
        doc,
        fig4,
        "Hình 4. Figure 4 trong contribution1.pdf: cơ chế edge-aware attention, trong đó omega được chiếu thành một bias phụ thêm vào attention score.",
        Inches(5.8),
    )
    add_captioned_picture(
        doc,
        fig5,
        "Hình 5. Figure 5 trong contribution1.pdf: sau nhiều lớp GAT, các biểu diễn node được pooling thành vector đồ thị cấp tài liệu.",
        Inches(5.8),
    )
    add_display_equation(doc, "e_uv^(k) = <W_q^(k) h_u, W_k^(k) h_v> / √d_k + Proj(ω_uv)")
    add_display_equation(doc, "α_uv^(k) = exp(e_uv^(k)) / Σ_{w ∈ N(v)} exp(e_wv^(k))")
    add_display_equation(doc, "h_v^(l+1) = W_o [ ||_k Σ_{u ∈ N(v)} α_uv^(k) · ω_uv · W_v^(k) h_u^(l) ] + h_v^(l)")

    doc.add_heading("4.3. Constraint Score và Joint Scorer: KG được đưa vào retrieve như thế nào", level=2)
    add_paragraphs(
        doc,
        [
            "Graph embedding không phải là đầu ra duy nhất của KG. Song song với nhánh GAT, hệ thống còn tính một Constraint Score để đo mức độ nhất quán kế toán của bảng ứng viên. Trong constraint_score.py, điểm này được xây bằng cách lấy residual |ω · v_u - v_v| trên từng accounting edge, chuẩn hóa theo max(|v_v|, ε) rồi đưa qua hàm mũ âm.",
            "Điểm then chốt là Constraint Score đi vào Joint Scorer cùng với text similarity và entity score. Nghĩa là KG không trực tiếp thay thế semantic retrieval, mà đóng vai trò làm một nguồn tín hiệu cấu trúc bổ sung trong reranking. Đây là lý do ý tưởng hiện tại khả thi: ta không cần thay toàn bộ retriever, nhưng vẫn tăng khả năng tách đúng tài liệu khỏi các distractor cùng chủ đề.",
            "Mặt khác, Constraint Score cũng cho thấy rõ giới hạn hiện tại. Vì nhiều đẳng thức nhiều toán hạng đang bị giản lược thành các cạnh cặp đôi, điểm số này chỉ là kiểm tra mềm theo cạnh, chưa phải kiểm tra phương trình trung thành theo nghĩa reasoning. Về mặt nghiên cứu, đây là chỗ nên được nâng cấp trong giai đoạn tiếp theo.",
        ],
    )
    add_captioned_picture(
        doc,
        fig6,
        "Hình 6. Figure 6 trong contribution1.pdf: Joint Scorer hợp nhất ba tín hiệu s_text, s_ent và CS(G_D) bằng các trọng số học được α, β, γ.",
        Inches(5.7),
    )
    add_display_equation(
        doc,
        "CS(G_D) = (1 / |E_c|) Σ_(u,v,ω) exp( - |ω · v_u - v_v| / max(|v_v|, ε) )"
    )

    doc.add_heading("4.4. CACL và CHAP: phần huấn luyện đã được đề xuất và đã có code", level=2)
    add_paragraphs(
        doc,
        [
            "Nếu GSR là đóng góp ở phía inference, thì CACL là đóng góp ở phía training. Ý tưởng cơ bản của CACL là không để retriever học từ những negative quá dễ, mà bắt nó phân biệt với các negative nhìn rất giống bản gốc nhưng bị phá đúng một ràng buộc quan trọng. CHAP là cơ chế tạo ra các negative như vậy.",
            "Trong chap.py, ba kiểu negative đã được hiện thực hóa tương đối rõ: CHAP-A sửa một ô dữ liệu để phá đẳng thức cộng; CHAP-S làm sai scale của giá trị; CHAP-E hoán đổi company hoặc year trong metadata. Đây là một thiết kế rất sát với lỗi thực tế của retrieval tài chính, nơi mô hình thường không sai ở mức chủ đề, mà sai ở mức rất tinh vi như đúng công ty nhưng lệch năm hoặc đúng bảng nhưng sai scale.",
            "CACL cũng không chỉ dừng ở negative sampler. train.py và trainer.py đã thể hiện rõ một lịch huấn luyện ba giai đoạn: Identity Pretraining để dạy mô hình phân biệt metadata, Structural Pretraining để hiệu chỉnh thành phần graph/constraint, và Joint CACL để tối ưu toàn hệ với CHAP negatives. Vì vậy, có thể khẳng định rằng CACL đã tồn tại ở mức thiết kế và code, dù chưa được đóng kín thành một pipeline benchmark hoàn chỉnh như phần retrieval.",
        ],
    )
    add_captioned_picture(
        doc,
        fig7,
        "Hình 7. Figure 7 trong contribution1.pdf: CHAP sinh hard negatives theo ba kiểu A, S, E; mỗi kiểu chỉ phá vỡ đúng một bất biến tài chính hoặc metadata.",
        Inches(5.4),
    )
    add_captioned_picture(
        doc,
        fig8,
        "Hình 8. Figure 8 trong contribution1.pdf: hàm mục tiêu CACL kết hợp triplet loss với penalty cho vi phạm ràng buộc.",
        Inches(5.5),
    )
    add_captioned_picture(
        doc,
        fig9,
        "Hình 9. Figure 9 trong contribution1.pdf: lịch huấn luyện ba giai đoạn Identity → Structural → Joint CACL.",
        Inches(5.6),
    )
    add_display_equation(doc, "L_CACL = L_triplet + λ · L_constraint")

    stage_table = doc.add_table(rows=1, cols=4)
    stage_table.style = "Table Grid"
    stage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Giai đoạn", "Mục tiêu", "Đã thể hiện ở đâu", "Đánh giá tiến độ"]):
        set_cell_text(stage_table.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(stage_table.rows[0].cells[i], "D9D2E9")
    set_repeat_table_header(stage_table.rows[0])
    stage_rows = [
        ("Stage 1: Identity", "Học phân biệt company, year, sector", "train.py; joint_scorer.py; compute_entity_score", "Có nền tảng tốt nhưng metadata còn khai thác hẹp"),
        ("Stage 2: Structural", "Calibrate graph encoder và constraint signal", "gat_layer.py; constraint_score.py", "Đã có cơ chế rõ ở mức prototype"),
        ("Stage 3: Joint CACL", "Tối ưu toàn hệ với CHAP negatives", "negative_sampler/chap.py; trainer.py", "Đã có khung train, chưa chứng minh end-to-end trên reasoning"),
    ]
    for row_data in stage_rows:
        cells = stage_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("4.5. Chốt lại những gì đã thực sự làm được", level=2)
    add_paragraphs(
        doc,
        [
            "Từ góc nhìn tiến độ, có thể chốt khách quan rằng GSR đã được triển khai rõ nhất ở mức inference retrieval. CACL đã được triển khai ở mức thiết kế huấn luyện và negative generation, nhưng dấu nối giữa training objective và benchmark inference hiện vẫn chưa kín như một hệ hoàn chỉnh. Constraint KG đã có giá trị thực tiễn ở vai trò structural reranking, song chưa phải một knowledge substrate đủ mạnh cho reasoning số học đa bước.",
            "Nói cách khác, đề tài hiện đã có đóng góp thực và có thể báo cáo rõ với GVHD ở hai phương diện: một là đã xây được retrieval có cấu trúc thay vì retrieval text-only; hai là đã chỉ ra đúng nơi mà retrieval hiện tại dừng lại, từ đó mở ra định hướng mới có chiều sâu hơn là Financial Evidence Graph cho cả retrieval và reasoning.",
        ],
    )

    doc.add_heading("5. Đánh giá khách quan về khả năng dùng KG cho reasoning số học", level=1)
    add_paragraphs(
        doc,
        [
            "Ý tưởng dùng KG đã xây cho retrieval để đẩy tiếp sang reasoning là khả thi, nhưng không thể dùng nguyên trạng. Tính khả thi nằm ở chỗ KG hiện tại đã chứng minh được một điều rất quan trọng: biểu diễn bảng tài chính bằng cấu trúc có hướng thực sự giúp hệ thống nhận biết tốt hơn sự khác biệt giữa các tài liệu bề ngoài rất giống nhau. Đây là tiền đề rất mạnh để đi tiếp.",
            "Vấn đề nằm ở chỗ biểu diễn hiện tại của KG còn quá thiên về reranking tài liệu. Đồ thị hiện chưa lưu provenance đủ rõ, chưa liên kết đồng thời text - table - footnote, chưa chuẩn hóa mạnh unit/scale, chưa có node phương trình hay node phép toán, và chưa hỗ trợ truy vết từ một đáp án ngược trở lại tập toán hạng sinh ra nó. Nếu dùng nguyên KG hiện nay để reasoning, mô hình sẽ vẫn phải suy luận trên context khá thô và dễ trộn nhiễu.",
            "Bởi vậy, kết luận khách quan là ý tưởng này nên triển khai, nhưng phải triển khai theo hướng nâng cấp KG thành một evidence graph thống nhất, chứ không nên chỉ nối thẳng retriever hiện có với một module suy luận số học ở đầu sau. Nếu làm ngắn mạch như vậy, lợi ích của graph sẽ bị giới hạn ở retrieval và không đi vào chỗ khó nhất là chọn đúng toán hạng dưới top-k nhiễu.",
        ],
    )

    gap_table = doc.add_table(rows=1, cols=3)
    gap_table.style = "Table Grid"
    gap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Vấn đề hiện tại", "Hệ quả", "Hướng nâng cấp cần làm"]):
        set_cell_text(gap_table.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(gap_table.rows[0].cells[i], "F4CCCC")
    set_repeat_table_header(gap_table.rows[0])
    gap_rows = [
        ("Metadata mới dừng ở company, year, sector", "Reranking có ích nhưng chưa đủ mạnh để kiểm soát reasoning", "Nâng thành ontology truy xuất với quarter, statement type, unit, scale, aliases, provenance"),
        ("Constraint Score kiểm tra theo cạnh cặp đôi", "Không đủ trung thành cho phương trình nhiều toán hạng", "Đưa equation node hoặc hyperedge vào Financial Evidence Graph"),
        ("Retrieval còn thiên về document level", "Reasoning phải đối mặt top-k còn nhiều nhiễu", "Phân cấp retrieval: document → section/table → evidence atom"),
        ("KG chưa nối với text và footnote", "Dễ bỏ sót ngoại lệ hoặc điều kiện nằm ngoài bảng", "Thống nhất bảng, đoạn văn và footnote trên cùng đồ thị bằng chứng"),
    ]
    for row_data in gap_rows:
        cells = gap_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("6. Định hướng chốt: Financial Evidence Graph dùng chung cho retrieval và reasoning", level=1)
    add_paragraphs(
        doc,
        [
            "Định hướng được chốt cho giai đoạn tiếp theo là xây dựng một Financial Evidence Graph có kiểu, dùng chung cho cả retrieval và reasoning. Trong kiến trúc này, KG không còn là một mảnh phụ trợ đặt cạnh retriever, mà trở thành lớp biểu diễn trung tâm của toàn bộ pipeline. Mỗi câu hỏi trước hết đi qua query parser để rút ra company, period, metric và kiểu phép toán; sau đó hệ thống truy xuất phân cấp để tạo local evidence graph quanh câu hỏi; cuối cùng planner, executor và verifier đều làm việc trên lớp graph này.",
            "Điểm then chốt của định hướng mới là metadata phải được nâng thành ontology truy xuất. Thay vì chỉ có company, year, sector, evidence atom cần được gắn rõ statement type, quarter, đơn vị, scale, currency, row path, column path và nguồn gốc từ document nào. Khi đó metadata không chỉ góp điểm ở scorer, mà còn tham gia pre-filtering, chunk embedding, hard-negative mining, grounding và verifier.",
            "Ở phía reasoning, local evidence graph cần hỗ trợ ba chức năng rõ ràng. Thứ nhất là chọn đúng toán hạng từ top-k còn nhiễu. Thứ hai là sinh một chương trình ngắn dạng DSL hoặc Python để thực thi phép tính. Thứ ba là kiểm tra ngược kết quả bằng cách xác minh company, year, đơn vị, scale, đồng thời kiểm tra xem chương trình vừa chạy có thực sự grounded trên các node bằng chứng hay không.",
        ],
    )
    add_captioned_picture(
        doc,
        fig_path,
        "Hình 10. Kiến trúc đề xuất cho giai đoạn tiếp theo: phía trên là retrieval GSR-CACL hiện tại; phía dưới là kiến trúc Financial Evidence Graph nhằm thống nhất retrieval, grounding, reasoning và verification.",
        Inches(6.2),
    )

    graph_table = doc.add_table(rows=1, cols=3)
    graph_table.style = "Table Grid"
    graph_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Thành phần của Financial Evidence Graph", "Vai trò", "Lợi ích trực tiếp"]):
        set_cell_text(graph_table.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(graph_table.rows[0].cells[i], "D9EAF7")
    set_repeat_table_header(graph_table.rows[0])
    graph_rows = [
        ("Document / Section / Table / Cell / Sentence / Footnote", "Tạo lớp bằng chứng đa hạt và có provenance", "Cho phép retrieval không dừng ở cấp tài liệu"),
        ("Company / Period / Statement / Unit / Scale", "Ontology metadata cho retrieval và verification", "Giảm nhầm lẫn đúng công ty nhưng sai năm, đúng năm nhưng sai statement"),
        ("MetricConcept / Equation / Operation node", "Biểu diễn rõ các quan hệ số học nhiều toán hạng", "Executor và verifier có thể reasoning trực tiếp trên graph"),
        ("Support / refers_to / derived_from / same_metric_as", "Nối table với văn bản và footnote", "Tăng khả năng grounding và giải thích kết quả"),
    ]
    for row_data in graph_rows:
        cells = graph_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("7. Kế hoạch triển khai tiếp theo", level=1)
    add_paragraphs(
        doc,
        [
            "Kế hoạch triển khai nên được chia theo lớp năng lực, không nên nhảy thẳng sang RL cho reasoning khi retrieval và evidence grounding còn chưa ổn định. Thứ tự hợp lý là: củng cố retrieval có metadata-aware filtering; xây Financial Evidence Graph; hiện thực planner/executor/verifier; sau đó mới tiến hành huấn luyện preference optimization và reinforcement learning trên reasoning trace.",
            "Đặc biệt, phần học tăng cường cần được xem là giai đoạn tối ưu cuối, không phải điểm khởi đầu. Nếu verifier và executor chưa đủ tin cậy thì reward của GRPO sẽ nhiễu, dẫn đến học sai mục tiêu. Vì vậy, giá trị lớn nhất của giai đoạn hiện tại là chốt chuẩn kiến trúc và chia việc theo phụ thuộc kỹ thuật một cách chặt chẽ.",
        ],
    )

    roadmap = doc.add_table(rows=1, cols=4)
    roadmap.style = "Table Grid"
    roadmap.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Giai đoạn", "Mục tiêu", "Công việc chính", "Chỉ số / đầu ra cần có"]):
        set_cell_text(roadmap.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(roadmap.rows[0].cells[i], "EAD1DC")
    set_repeat_table_header(roadmap.rows[0])
    roadmap_rows = [
        ("Giai đoạn 1", "Củng cố retrieval", "Mở rộng metadata schema; thêm pre-filtering; table/section reranking; đánh giá ablation giữa text-only và GSR", "MRR/Recall/NDCG tăng rõ; có phân tích lỗi theo company/year/statement"),
        ("Giai đoạn 2", "Xây Financial Evidence Graph", "Thiết kế node/edge schema; lưu provenance; nối text, table, footnote; thêm equation node", "Có graph builder mới và case study trực quan hóa"),
        ("Giai đoạn 3", "Grounding và reasoning substrate", "Query parser; operand grounding; DSL/Python executor; verifier kiểm tra unit/year/scale", "Đo được operand accuracy, execution accuracy và grounding accuracy"),
        ("Giai đoạn 4", "Huấn luyện reasoning", "SFT cho chương trình suy luận; Step-DPO cho trace; GRPO/RLVR với reward verifiable", "Cải thiện answer accuracy dưới top-k nhiễu"),
        ("Giai đoạn 5", "Đánh giá và viết bài", "Thiết kế benchmark end-to-end; robustness test; ablation KG/metadata/verifier", "Bộ kết quả đủ chặt cho báo cáo và bài báo hội nghị"),
    ]
    for row_data in roadmap_rows:
        cells = roadmap.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("7.1. Định hướng lựa chọn thuật toán huấn luyện", level=2)
    add_paragraphs(
        doc,
        [
            "Đối với retrieval, CACL và CHAP vẫn nên được giữ lại vì chúng đúng với bản chất negative của tài liệu tài chính. Tuy nhiên, ở phần reasoning không nên kỳ vọng CACL một mình giải quyết bài toán. Cách hợp lý hơn là dùng Supervised Fine-Tuning để dạy dạng chương trình suy luận chuẩn, sau đó dùng Step-DPO để ưu tiên trace grounded hơn trace bị nhiễu, và chỉ khi verifier đủ tốt mới chuyển sang GRPO hoặc một biến thể RLVR để tối ưu cuối.",
            "Trong ba thuật toán preference/RL thường được cân nhắc là DPO, ORPO và GRPO, ORPO phù hợp hơn như baseline tính toán rẻ; DPO hoặc Step-DPO phù hợp hơn cho giai đoạn làm sạch reasoning trace; còn GRPO phù hợp nhất cho bước tối ưu cuối cùng vì bài toán tài chính có reward kiểm chứng được từ executor và verifier. Cách phân vai như vậy nhất quán hơn với cấu trúc kỹ thuật của đề tài so với việc chọn một thuật toán duy nhất cho toàn bộ pipeline.",
        ],
    )

    algo_table = doc.add_table(rows=1, cols=4)
    algo_table.style = "Table Grid"
    algo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Thuật toán", "Nên dùng ở đâu", "Ưu điểm", "Lưu ý"]):
        set_cell_text(algo_table.rows[0].cells[i], head, bold=True, font_size=10)
        shade_cell(algo_table.rows[0].cells[i], "FFF2CC")
    set_repeat_table_header(algo_table.rows[0])
    algo_rows = [
        ("CACL + CHAP", "Retrieval training", "Học phân biệt negative rất gần nhưng sai logic tài chính", "Không thay thế được reasoning optimizer"),
        ("SFT", "Reasoning bootstrap", "Ổn định, dễ kiểm soát chương trình vàng", "Cần dữ liệu reasoning trace có chất lượng"),
        ("Step-DPO / DPO", "Trace preference", "Ưu tiên suy luận grounded hơn suy luận nhiễu", "Cần chuẩn bị cặp trace tốt/xấu đủ rõ"),
        ("GRPO / RLVR", "Tối ưu reasoning cuối", "Tận dụng reward kiểm chứng được từ executor", "Chỉ nên dùng khi verifier đủ đáng tin"),
    ]
    for row_data in algo_rows:
        cells = algo_table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("8. Kết luận", level=1)
    add_paragraphs(
        doc,
        [
            "Nhìn từ trạng thái hiện tại, đề tài đã có phần lõi đủ mạnh để báo cáo một cách tự tin: GSR, CACL và Constraint KG đều đã xuất hiện dưới dạng đóng góp kỹ thuật có mã nguồn tương ứng. Phần retrieval hiện tại không còn là ý tưởng mơ hồ, mà là một prototype nghiên cứu thực sự có thể benchmark và phân tích.",
            "Điểm cần chốt với GVHD là định hướng phát triển tiếp theo phải chuyển từ KG cho reranking sang Financial Evidence Graph dùng chung cho retrieval và reasoning. Đây là bước nâng cấp làm cho đồ thị tri thức không chỉ giúp tìm đúng tài liệu, mà còn giúp neo đúng toán hạng, thực thi đúng phép toán và kiểm tra được đáp án cuối cùng. Nếu bám chặt theo lộ trình này, đề tài có nền tảng tốt để tiến tới một đóng góp học thuật sâu hơn và thuyết phục hơn nhiều so với việc chỉ tiếp tục tối ưu MRR ở tầng retrieval.",
        ],
    )

    doc.add_heading("Tài liệu tham khảo chính", level=1)
    refs = [
        "Strich, J., Isgorur, E. K., Trescher, M., Biemann, C., & Semmann, M. T²-RAGBench: Text-and-Table Benchmark for Evaluating Retrieval-Augmented Generation. arXiv:2506.12071, 2025; accepted to EACL 2026.",
        "Chen, Z. et al. FinQA: A Dataset of Numerical Reasoning over Financial Data. EMNLP 2021.",
        "Zhu, F. et al. TAT-QA: A Question Answering Benchmark on a Hybrid of Tabular and Textual Content in Finance. ACL 2021.",
        "Chen, Z. et al. ConvFinQA: Exploring the Chain of Numerical Reasoning in Conversational Finance Question Answering. EMNLP 2022.",
        "Dang, Q.-V., Nguyen, N.-S.-A., & Vo, T.-B.-D. HierFinRAG—Hierarchical Multimodal RAG for Financial Document Understanding. Informatics 13(2):30, 2026.",
        "FT-RAG: A Fine-grained Retrieval-Augmented Generation Framework for Financial Analysis. arXiv:2605.01495, 2026.",
        "Rafailov, R. et al. Direct Preference Optimization: Your Language Model is Secretly a Reward Model. arXiv:2305.18290, 2023.",
        "Hong, J. et al. ORPO: Monolithic Preference Optimization without Reference Model. EMNLP 2024.",
        "Shao, Z. et al. DeepSeekMath: Pushing the Limits of Mathematical Reasoning in Open Language Models. arXiv:2402.03300, 2024.",
        "Tài liệu nội bộ: contribution1.pdf, bao_cao_tien_do_nghien_cuu_gvhd_2026-06-13.md và mã nguồn trong thư mục NLP/ours/source.",
    ]
    add_reference_list(doc, refs)
    doc.save(docx_path)


def main() -> None:
    build_diagram(FIG_PATH)
    build_progress_report(OUT_DOCX, FIG_PATH)
    print(f"Generated: {OUT_DOCX}")
    print(f"Generated: {FIG_PATH}")


if __name__ == "__main__":
    main()
