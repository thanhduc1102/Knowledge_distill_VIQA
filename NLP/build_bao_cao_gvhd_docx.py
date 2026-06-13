from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "figures"
OUT_DOCX = ROOT / "bao_cao_chien_luoc_kg_retrieval_reasoning_gvhd_2026-06-13.docx"
FIG_PATH = FIG_DIR / "kien_truc_retrieval_reasoning_2026-06-13.png"


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    font_path = Path("C:/Windows/Fonts") / font_name
    return ImageFont.truetype(str(font_path), size=size)


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill):
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=4, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x = x0 + (x1 - x0 - tw) / 2
    y = y0 + (y1 - y0 - th) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=4, align="center")


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, fill, outline, title_fill=(20, 20, 20)):
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, title, get_font(24, bold=True), title_fill)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(60, 60, 60), width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        p1 = (ex - 18 * sign, ey - 8)
        p2 = (ex - 18 * sign, ey + 8)
    else:
        sign = 1 if dy >= 0 else -1
        p1 = (ex - 8, ey - 18 * sign)
        p2 = (ex + 8, ey - 18 * sign)
    draw.polygon([end, p1, p2], fill=color)


def build_diagram(out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2200, 1500), (250, 250, 252))
    draw = ImageDraw.Draw(img)

    title_font = get_font(40, bold=True)
    body_font = get_font(22)
    section_font = get_font(28, bold=True)

    draw.text((60, 40), "Kiến trúc mô hình: hiện trạng retrieval và hướng phát triển toàn trình", font=title_font, fill=(10, 40, 95))

    # Panel A
    panel_a = (50, 130, 2150, 700)
    panel_b = (50, 780, 2150, 1425)
    draw.rounded_rectangle(panel_a, radius=20, outline=(20, 82, 147), width=4, fill=(240, 247, 255))
    draw.rounded_rectangle(panel_b, radius=20, outline=(0, 110, 82), width=4, fill=(241, 252, 247))
    draw.text((80, 150), "A. Kiến trúc retrieval hiện tại theo GSR-CACL", font=section_font, fill=(20, 82, 147))
    draw.text((80, 800), "B. Kiến trúc đề xuất cho retrieval + reasoning toàn trình", font=section_font, fill=(0, 110, 82))

    a_boxes = [
        ((100, 260, 360, 390), "Câu hỏi +\nmetadata truy vấn", (255, 255, 255), (20, 82, 147)),
        ((430, 260, 710, 390), "FAISS lấy tập\nứng viên tài liệu", (255, 255, 255), (20, 82, 147)),
        ((780, 260, 1070, 390), "Trích bảng markdown\nđầu tiên", (255, 255, 255), (20, 82, 147)),
        ((1140, 210, 1450, 340), "Template matching\nvà xây Constraint KG", (255, 255, 255), (20, 82, 147)),
        ((1140, 400, 1450, 530), "Constraint score", (255, 255, 255), (20, 82, 147)),
        ((1520, 210, 1840, 340), "Edge-aware GAT\n→ kg_embed", (255, 255, 255), (20, 82, 147)),
        ((1520, 400, 1840, 530), "Entity score\n(company/year/sector)", (255, 255, 255), (20, 82, 147)),
        ((1900, 260, 2110, 390), "Joint scorer\n→ Top-k", (255, 255, 255), (20, 82, 147)),
    ]
    for box, text, fill, outline in a_boxes:
        draw_box(draw, box, text, fill, outline)

    draw_arrow(draw, (360, 325), (430, 325), color=(20, 82, 147))
    draw_arrow(draw, (710, 325), (780, 325), color=(20, 82, 147))
    draw_arrow(draw, (1070, 325), (1140, 275), color=(20, 82, 147))
    draw_arrow(draw, (1070, 325), (1140, 465), color=(20, 82, 147))
    draw_arrow(draw, (1450, 275), (1520, 275), color=(20, 82, 147))
    draw_arrow(draw, (1450, 465), (1900, 355), color=(20, 82, 147))
    draw_arrow(draw, (1840, 275), (1900, 325), color=(20, 82, 147))
    draw_arrow(draw, (1840, 465), (1900, 355), color=(20, 82, 147))

    draw.text(
        (120, 565),
        "Điểm nghẽn chính của kiến trúc hiện tại: retrieval mới mạnh ở cấp tài liệu, KG vẫn thiên về cấu trúc reranking\n"
        "và chưa đóng vai trò đầy đủ trong việc định vị toán hạng, lập phương trình hay kiểm chứng suy luận số học.",
        font=body_font,
        fill=(40, 40, 40),
        spacing=6,
    )

    # Panel B
    b_boxes = [
        ((95, 915, 315, 1040), "Câu hỏi", (255, 255, 255), (0, 110, 82)),
        ((360, 860, 620, 995), "Phân tích truy vấn:\ncompany / time /\nmetric / operation", (255, 255, 255), (0, 110, 82)),
        ((360, 1035, 620, 1170), "Ontology metadata:\nstatement, quarter,\nunit, scale, aliases", (255, 255, 255), (0, 110, 82)),
        ((680, 915, 960, 1040), "Retrieval phân cấp:\ndocument → table/\nsection → atom", (255, 255, 255), (0, 110, 82)),
        ((1015, 835, 1320, 970), "Financial Evidence Graph:\ncell, text, footnote,\nequation, provenance", (255, 255, 255), (0, 110, 82)),
        ((1015, 1080, 1320, 1215), "Evidence triage:\nlọc top-3 nhiễu,\nneo đúng toán hạng", (255, 255, 255), (0, 110, 82)),
        ((1380, 835, 1685, 970), "Reasoning planner:\nchọn phép toán,\nsinh DSL/Python", (255, 255, 255), (0, 110, 82)),
        ((1380, 1080, 1685, 1215), "Executor:\nthực thi số học", (255, 255, 255), (0, 110, 82)),
        ((1750, 915, 2045, 1040), "Verifier:\nkiểm tra grounding,\nunit, year,\nconstraint", (255, 255, 255), (0, 110, 82)),
    ]
    for box, text, fill, outline in b_boxes:
        draw_box(draw, box, text, fill, outline)

    draw_arrow(draw, (315, 975), (360, 925), color=(0, 110, 82))
    draw_arrow(draw, (315, 975), (360, 1100), color=(0, 110, 82))
    draw_arrow(draw, (620, 925), (680, 965), color=(0, 110, 82))
    draw_arrow(draw, (620, 1100), (680, 990), color=(0, 110, 82))
    draw_arrow(draw, (960, 975), (1015, 900), color=(0, 110, 82))
    draw_arrow(draw, (960, 975), (1015, 1148), color=(0, 110, 82))
    draw_arrow(draw, (1320, 900), (1380, 900), color=(0, 110, 82))
    draw_arrow(draw, (1320, 1148), (1380, 1148), color=(0, 110, 82))
    draw_arrow(draw, (1685, 900), (1750, 960), color=(0, 110, 82))
    draw_arrow(draw, (1685, 1148), (1750, 995), color=(0, 110, 82))

    draw.text(
        (120, 1260),
        "Ý tưởng chốt: biến KG thành Financial Evidence Graph dùng chung cho retrieval và reasoning, để mô hình\n"
        "không chỉ chọn đúng tài liệu mà còn chọn đúng toán hạng, thực thi đúng chương trình và tự kiểm chứng kết quả.",
        font=body_font,
        fill=(35, 45, 35),
        spacing=6,
    )

    img.save(out_path)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, RGBColor(18, 60, 110)),
        ("Heading 1", 16, RGBColor(18, 60, 110)),
        ("Heading 2", 13, RGBColor(32, 32, 32)),
        ("Heading 3", 12, RGBColor(64, 64, 64)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Caption" not in doc.styles:
        caption_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        caption_style.font.size = Pt(10.5)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(90, 90, 90)
    else:
        caption_style = doc.styles["Caption"]
        caption_style.font.name = "Times New Roman"
        caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        caption_style.font.size = Pt(10.5)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(90, 90, 90)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("BÁO CÁO CHIẾN LƯỢC RETRIEVAL VÀ SUY LUẬN TÀI CHÍNH")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer.add_run("Trang ")
    r1.font.name = "Times New Roman"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r1.font.size = Pt(9)
    add_page_number(footer)


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)


def add_reference_list(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.add_run(f"[{idx}] {item}")


def build_report(docx_path: Path, fig_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    configure_styles(doc)
    add_header_footer(doc)

    core = doc.core_properties
    core.title = "Báo cáo chiến lược truy xuất và suy luận tài chính dựa trên KG"
    core.author = "OpenAI Codex"
    core.subject = "Báo cáo nghiên cứu gửi GVHD"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.add_run("BÁO CÁO CHIẾN LƯỢC\nTRUY XUẤT VÀ SUY LUẬN TÀI CHÍNH DỰA TRÊN KG").bold = True
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(18, 60, 110)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(20)
    sub.add_run("Tổng hợp benchmark, hiện trạng triển khai, phân tích kỹ thuật, định hướng nghiên cứu\nvà kế hoạch phát triển toàn trình cho bài toán tài chính").font.size = Pt(13)
    for run in sub.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(140)
    info.add_run("Ngày cập nhật: 13/06/2026\nTài liệu phục vụ báo cáo với giáo viên hướng dẫn").font.size = Pt(12)
    for run in info.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_page_break()

    doc.add_heading("1. Tóm tắt điều hành", level=1)
    add_paragraphs(
        doc,
        [
            "Mục tiêu của báo cáo này là tổng hợp một cách hệ thống toàn bộ những gì đã được xây dựng trong repo hiện tại, đối chiếu với bài toán nghiên cứu thực sự cần giải quyết, và từ đó chốt lại một hướng triển khai đủ mạnh cho giai đoạn tiếp theo. Trọng tâm của đề tài không còn nên được hiểu là bài toán tăng điểm retrieval đơn thuần, mà là bài toán truy xuất đúng bằng chứng, định vị đúng toán hạng và thực hiện suy luận số học đáng tin cậy trong bối cảnh tài liệu tài chính dài, nhiều bảng, nhiều kỳ báo cáo và có nhiễu retrieval.",
            "Qua quá trình khảo sát benchmark, đọc tài liệu liên quan, phân tích nội dung trong contribution1.pdf và đối chiếu trực tiếp với code của hệ thống GSR-CACL hiện tại, có thể khẳng định rằng repo đã có một nền kỹ thuật tốt ở mức retrieval có cấu trúc. Tuy nhiên, hệ thống hiện tại vẫn chủ yếu là một prototype retrieval-centric, trong đó KG mới đóng vai trò hỗ trợ reranking tài liệu chứ chưa thực sự trở thành hạ tầng bằng chứng dùng chung cho cả retrieval và reasoning.",
            "Kết luận chiến lược quan trọng nhất của báo cáo là cần chuyển trọng tâm nghiên cứu từ Constraint KG cho retrieval sang Financial Evidence Graph dùng chung cho retrieval và reasoning. Trong hướng mới này, metadata không còn là một bonus yếu cho reranker, mà phải được nâng thành ontology truy xuất có kiểu; còn reasoning không nên chạy trực tiếp trên top-k context, mà phải vận hành trên một đồ thị bằng chứng cục bộ đã được triage, grounding và kiểm chứng.",
        ],
    )

    doc.add_heading("2. Bối cảnh, benchmark và bản chất bài toán", level=1)
    add_paragraphs(
        doc,
        [
            "Các hệ thống hỏi đáp tài chính hiện đại gặp hai khó khăn bản chất. Thứ nhất, thông tin không nằm trong một đoạn văn ngắn mà trải dài trên nhiều báo cáo, nhiều bảng và nhiều footnote. Thứ hai, đáp án cuối cùng thường không phải là một fact có sẵn mà là kết quả của một quá trình định vị toán hạng, chuẩn hóa đơn vị và thực hiện phép tính nhiều bước. Chính vì vậy, các thiết lập QA giả định context vàng có sẵn thường không phản ánh đúng độ khó thực sự của bài toán.",
            "Trong bối cảnh đó, T2-RAGBench là benchmark phù hợp nhất với hướng đề tài vì nó đặt retrieval thành yêu cầu bắt buộc trước reasoning. Thay vì đánh giá trực tiếp mô hình sinh đáp án trên context đã biết đúng, benchmark này buộc hệ thống phải truy xuất từ corpus tài chính hỗn hợp gồm text và bảng rồi mới reasoning. Đây là setting gần thực tế hơn và cũng là lý do vì sao repo hiện tại bám vào T2-RAGBench là một lựa chọn hợp lý.",
            "Tuy nhiên, đánh giá chỉ bằng MRR@3 hay Recall@k là chưa đủ. Trong bài toán tài chính, hoàn toàn có thể xảy ra tình huống top-3 có một context đúng nhưng hai context còn lại gây nhiễu mạnh ở mức toán hạng. Khi đó, retrieval nhìn có vẻ tốt nhưng reasoning vẫn thất bại. Do đó, benchmark cần được nâng từ đánh giá tài liệu sang đánh giá nhiều tầng, bao gồm retrieval, evidence grounding, reasoning và robustness dưới nhiễu.",
        ],
    )

    doc.add_heading("2.1. Vai trò của các bộ dữ liệu trong lộ trình nghiên cứu", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Các bộ dữ liệu hiện có không đóng vai trò trùng lặp, mà mỗi bộ đóng góp một góc nhìn khác nhau vào bài toán. Bảng dưới đây tóm tắt vai trò của từng benchmark trong lộ trình đề tài.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Bộ dữ liệu", "Vai trò chính", "Đặc trưng nổi bật", "Mục tiêu sử dụng trong đề tài"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9EAF7")
    set_repeat_table_header(table.rows[0])

    dataset_rows = [
        ("FinQA", "Suy luận tài chính có chương trình vàng", "Nhiều câu hỏi cần phép tính nhiều bước trên bảng và văn bản", "Huấn luyện và đánh giá executor/program reasoning"),
        ("ConvFinQA", "Suy luận hội thoại tài chính", "Phụ thuộc liên lượt hỏi, reasoning dài hơn", "Đánh giá khả năng bám ngữ cảnh và reasoning nhiều bước"),
        ("TAT-QA", "Hybrid table-text QA", "Kết hợp bằng chứng bảng và đoạn văn", "Đánh giá grounding liên phương thức"),
        ("DocFinQA", "Ngữ cảnh dài", "Tăng độ khó retrieval thực tế", "Stress-test retrieval và evidence triage"),
        ("T2-RAGBench", "Benchmark retrieval-first", "Không giả định oracle context", "Đánh giá retrieval và pipeline toàn trình"),
    ]
    for row in dataset_rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=10)

    doc.add_heading("3. Hiện trạng triển khai trong repo", level=1)
    add_paragraphs(
        doc,
        [
            "Qua phân tích mã nguồn, có thể xác định repo hiện tại gồm hai lớp chính. Lớp thứ nhất là các baseline retrieval trong thư mục baseline, bao gồm dense FAISS, hybrid BM25 cộng dense, HyDE và một số biến thể query rewriting. Lớp thứ hai là prototype đề xuất trong thư mục ours, nơi triển khai GSR-CACL với tham vọng đưa cấu trúc tài chính vào retrieval.",
            "Điều đáng lưu ý là contribution1.pdf không chỉ là phần trình bày ý tưởng, mà nhiều thành phần trong đó thực sự đã được hiện thực hóa ở mức prototype. Cụ thể, phần Graph-Structured Retrieval được phản ánh trong mô-đun xây KG theo template, Edge-aware GAT, constraint scoring và joint scorer; còn phần Constraint-Aware Contrastive Learning được phản ánh trong bộ sinh CHAP negatives, ba giai đoạn huấn luyện và loss kết hợp triplet với penalty vi phạm ràng buộc.",
            "Tuy vậy, việc hệ thống đã được cài đặt không đồng nghĩa rằng nó đã khép kín về mặt thực nghiệm. Qua đối chiếu code, có những điểm rất quan trọng: benchmark hiện tại phản ánh GSR inference rõ hơn là full GSR+CACL; constraint score hiện tại mới là tín hiệu cấu trúc xấp xỉ; và train/inference còn tồn tại lệch trong cách xây KG. Những phát hiện này rất quan trọng để tránh ngộ nhận về độ hoàn thiện của hệ thống.",
        ],
    )

    doc.add_heading("3.1. Bản đồ thành phần kỹ thuật hiện có", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Bảng sau tóm tắt các mô-đun quan trọng trong repo hiện tại, mục tiêu của từng mô-đun và đánh giá ngắn gọn về trạng thái sử dụng.")

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["Thành phần", "Đường dẫn chính", "Chức năng", "Đánh giá hiện trạng"]
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9EAD3")
    set_repeat_table_header(table2.rows[0])

    module_rows = [
        ("Dataset wrapper", "ours/source/src/gsr_cacl/datasets/wrappers.py", "Load T2-RAGBench, build corpus, query và metadata", "Đã dùng được, là nền tốt cho giai đoạn tiếp theo"),
        ("GSR retrieval", "ours/source/src/gsr_cacl/methods/gsr_retrieval.py", "FAISS candidate retrieval và rerank bằng joint scorer", "Đã có prototype thật sự ở mức retrieval"),
        ("KG builder", "ours/source/src/gsr_cacl/kg/builder.py", "Sinh Constraint KG từ bảng markdown theo template", "Hữu ích cho retrieval nhưng chưa đủ trung thành cho reasoning"),
        ("Template library", "ours/source/src/gsr_cacl/templates/library.py", "Thư viện template IFRS/GAAP", "Có thể tái sử dụng như lớp tri thức nền"),
        ("Edge-aware GAT", "ours/source/src/gsr_cacl/encoders/gat_encoder.py", "Mã hóa KG thành graph embedding", "Có giá trị ở mức structural reranking"),
        ("Constraint scoring", "ours/source/src/gsr_cacl/scoring/constraint_score.py", "Tạo điểm tuân thủ ràng buộc", "Cần thay bằng biểu diễn phương trình trung thành hơn"),
        ("Joint scorer", "ours/source/src/gsr_cacl/scoring/joint_scorer.py", "Kết hợp text, entity và constraint", "Thiết kế hợp lý cho retrieval, chưa đủ cho reasoning"),
        ("CACL training", "ours/source/src/gsr_cacl/train.py", "Huấn luyện ba giai đoạn với CHAP negatives", "Đã có khung tốt nhưng chưa nối kín vào benchmark inference"),
        ("Benchmark", "ours/source/src/gsr_cacl/benchmark_gsr.py", "Đánh giá MRR/Recall/NDCG", "Phù hợp cho retrieval, chưa phải benchmark toàn trình"),
    ]
    for row in module_rows:
        cells = table2.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    doc.add_heading("4. Phân tích kỹ thuật retrieval hiện tại dựa trên contribution1.pdf và mã nguồn", level=1)
    add_paragraphs(
        doc,
        [
            "Phần retrieval trong contribution1.pdf có giá trị đặc biệt vì nó không trình bày GSR như một graph dùng để trang trí, mà như một nỗ lực đưa cấu trúc và ràng buộc kế toán vào đúng pha truy xuất. Tinh thần cốt lõi của proposal này là dense retrieval thuần text thường thất bại trên tài liệu tài chính vì hai lý do: một là hiện tượng nhầm lẫn thực thể khi nhiều báo cáo của cùng công ty qua các năm có bề mặt ngôn ngữ rất giống nhau; hai là hiện tượng làm phẳng bảng khiến mô hình mất thông tin về cấu trúc số học bên trong.",
            "Theo kiến trúc được trình bày trong contribution1.pdf và đối chiếu với code hiện tại, GSR vận hành theo quy trình sau. Câu hỏi và metadata liên quan được dùng để truy xuất một tập ứng viên từ FAISS. Với từng tài liệu ứng viên, hệ trích bảng markdown, dùng template tài chính để chuẩn hóa tiêu đề và xây một Constraint KG, sau đó mã hóa KG này bằng Edge-aware GAT để thu được graph embedding. Song song, hệ tính một constraint score phản ánh mức độ phù hợp của cấu trúc bảng. Cuối cùng, joint scorer kết hợp tín hiệu text, entity và constraint để sắp xếp lại các ứng viên.",
            "Điều cần nhấn mạnh là retrieval ở đây vẫn là retrieval hai tầng. Tầng thứ nhất là dense retrieval để gom candidate; tầng thứ hai mới là reranking có cấu trúc. Nói cách khác, KG hiện chưa thay thế bộ truy xuất chính, mà đóng vai trò là một mô-đun tái chấm điểm ứng viên theo logic tài chính. Đây là một lựa chọn hợp lý cho giai đoạn đầu vì nó cho phép chèn tri thức miền mà không phải tái thiết kế hoàn toàn toàn bộ hệ truy xuất.",
        ],
    )

    doc.add_paragraph()
    doc.add_picture(str(fig_path), width=Inches(6.3))
    cap = doc.add_paragraph("Hình 1. Kiến trúc mô hình: bên trên là pipeline retrieval hiện tại theo tinh thần GSR-CACL; bên dưới là kiến trúc toàn trình được đề xuất cho giai đoạn tiếp theo.", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.1. Constraint KG: từ bảng tài chính đến cấu trúc có hướng", level=2)
    add_paragraphs(
        doc,
        [
            "Constraint KG hiện tại được xây từ bảng markdown bằng ba bước chính. Bước thứ nhất là parse bảng thành hàng và cột, đồng thời chuyển các chuỗi biểu diễn số thành giá trị chuẩn hóa, bao gồm các trường hợp có dấu ngoặc âm, đơn vị triệu, tỷ hoặc phần trăm. Bước thứ hai là chuẩn hóa header tài chính, ví dụ Sales được ánh xạ về Revenue, Net Profit được ánh xạ về Net Income, Tax Expense được ánh xạ về Income Tax. Bước thứ ba là match bảng với một thư viện template IFRS/GAAP để sinh các cạnh accounting hoặc fallback sang các cạnh positional nếu match yếu.",
            "Ý tưởng này có giá trị ở chỗ nó biến bảng từ dạng văn bản phẳng thành một cấu trúc có hướng, nơi mỗi ô là một node và mỗi quan hệ tài chính hoặc quan hệ vị trí là một cạnh. Với cách làm này, mô hình có thể phân biệt tốt hơn giữa những bảng tuy giống nhau về mặt từ vựng nhưng khác nhau về cấu trúc kế toán. Đây chính là lợi thế quan trọng của GSR so với dense retrieval chỉ nhìn vào chuỗi text.",
            "Tuy nhiên, qua phân tích kỹ thuật có thể thấy Constraint KG hiện tại vẫn chủ yếu phù hợp với retrieval hơn là reasoning. Nhiều phương trình kế toán nhiều toán hạng đang bị phân rã thành các cạnh cặp đôi đi vào cùng một node đích. Biểu diễn này tạo được tín hiệu cấu trúc, nhưng chưa đủ trung thành để kiểm chứng một phương trình hay để làm nền cho executor suy luận số học nhiều bước.",
        ],
    )

    doc.add_heading("4.2. Edge-aware GAT: cơ chế tạo graph embedding", level=2)
    add_paragraphs(
        doc,
        [
            "Edge-aware GAT trong hệ hiện tại là mô-đun học được quan trọng nhất ở phía KG. Mỗi node trước hết được mã hóa từ ba loại thông tin: nhãn header đã chuẩn hóa, đặc trưng số của giá trị trong ô và mã hóa vị trí theo hàng và cột. Từ đó, mạng GAT lan truyền thông tin dọc theo các cạnh của đồ thị để tạo ra embedding cho từng node và sau cùng gộp chúng thành một graph embedding cấp tài liệu.",
            "Điểm khác biệt quan trọng so với GAT thông thường là trọng số cạnh omega không chỉ tồn tại như một nhãn phụ mà đi trực tiếp vào cơ chế attention. Trong triển khai hiện tại, omega vừa tạo bias cho attention, vừa ảnh hưởng đến message passing giữa node nguồn và node đích. Điều này làm cho cạnh cộng, cạnh trừ và cạnh vị trí mang ý nghĩa khác nhau trong quá trình mã hóa. Nhờ vậy, KG không chỉ nói rằng hai ô có liên quan, mà còn nói chúng liên quan theo loại phép toán nào.",
            "Ở góc nhìn retrieval, điều này rất hữu ích vì graph embedding sẽ mang thêm dấu vết của cấu trúc bảng thay vì chỉ mang nghĩa từ vựng. Trong các tình huống có nhiều báo cáo gần giống nhau, sự khác biệt nhỏ về quan hệ hàng, cột, tổng và thành phần có thể là tín hiệu đủ mạnh để tách đúng tài liệu khỏi các distractor khó. Tuy nhiên, hiệu quả này vẫn phụ thuộc mạnh vào việc template matching có tốt hay không, bởi nếu đồ thị rơi vào fallback positional edges thì tín hiệu học được sẽ yếu hơn đáng kể.",
        ],
    )

    doc.add_heading("4.3. Constraint scoring: tín hiệu kiểm tra cấu trúc", level=2)
    add_paragraphs(
        doc,
        [
            "Bên cạnh graph embedding học được, hệ còn xây một tín hiệu thủ công gọi là constraint score. Về trực giác, đây là một chỉ báo cho biết một bảng ứng viên có vẻ nhất quán về mặt cấu trúc tài chính hay không. Tín hiệu này được tính bằng cách duyệt các cạnh accounting trong đồ thị, đo sai lệch giữa node nguồn và node đích theo omega, sau đó chuyển sai lệch đó qua một hàm mũ để thu về điểm trong khoảng từ 0 đến 1.",
            "Ưu điểm của constraint score là nó bổ sung một prior tài chính tương đối trực tiếp vào quá trình reranking. Trong nhiều trường hợp, hai tài liệu có thể rất giống nhau về semantic similarity, nhưng tài liệu đúng sẽ có cấu trúc số liệu hợp lý hơn. Constraint score giúp mô hình tận dụng được khác biệt này. Đây là một trong những điểm mạnh thực dụng nhất của kiến trúc retrieval hiện tại.",
            "Dù vậy, constraint score cũng là nơi bộc lộ rõ nhất giới hạn của thiết kế hiện tại. Vì các ràng buộc nhiều toán hạng đang được quy về các cạnh cặp đôi, điểm số sinh ra chỉ là một phép xấp xỉ, không phải kiểm chứng phương trình đúng nghĩa. Hệ quả là bảng đúng chưa chắc đạt điểm tối đa, còn bảng không match template mạnh đôi khi lại không bị phạt đủ rõ. Điều này củng cố nhận định rằng constraint score hiện tại có giá trị trong retrieval, nhưng không thể là bộ kiểm định chính cho reasoning.",
        ],
    )

    doc.add_heading("4.4. Joint scorer và vai trò thực của metadata", level=2)
    add_paragraphs(
        doc,
        [
            "Joint scorer là nơi ba nguồn tín hiệu được hợp nhất: text similarity, entity score và constraint signal. Đây là phần rất quan trọng vì nó cho thấy GSR không phủ nhận semantic retrieval, mà chỉ đặt semantic retrieval vào một khung scoring rộng hơn có tri thức miền. Nhờ cách hợp nhất này, hệ có thể ưu tiên các tài liệu vừa giống câu hỏi về nội dung, vừa đúng công ty, đúng năm và có cấu trúc tài chính hợp lý.",
            "Metadata hiện tại gồm company, year và sector. Ý tưởng ban đầu là dùng chúng như tín hiệu phân biệt thực thể, đặc biệt để khắc phục hiện tượng cùng công ty nhưng khác năm. Đây là một ý tưởng đúng, nhưng qua phân tích có thể thấy metadata vẫn chưa phát huy hết vai trò. Trong hệ hiện tại, metadata mới chủ yếu tham gia ở mức scoring bề mặt, chứ chưa thật sự đi vào biểu diễn chunk, vào hard-negative mining theo kiểu ontology, hay vào reasoning như một điều kiện tương thích của toán hạng.",
            "Điểm này đặc biệt quan trọng cho giai đoạn tiếp theo. Nếu coi metadata chỉ là bonus score, đóng góp của nó sẽ bị giới hạn. Nếu coi metadata là một ontology nhẹ của bài toán tài chính, nó có thể tham gia vào toàn bộ pipeline: điều kiện định tuyến candidate, cấu trúc chunk embedding, điều kiện triage evidence, ràng buộc chọn toán hạng và cả verifier hậu nghiệm. Đây là một hướng mở rất giàu tiềm năng mà hệ hiện tại mới chạm đến ở mức sơ khai.",
        ],
    )

    doc.add_heading("5. CACL và ý nghĩa của học tăng cường / tối ưu ưu tiên", level=1)
    add_paragraphs(
        doc,
        [
            "Trong contribution1.pdf, CACL được đưa ra như cơ chế huấn luyện nhằm làm cho retriever và scorer phân biệt được các tài liệu gần đúng nhưng sai về logic tài chính. Điểm hay của ý tưởng này nằm ở CHAP negatives: thay vì chỉ chọn negative khác chủ đề, hệ chủ động tạo các negative khó bằng cách làm sai một công thức, đổi scale hoặc đổi metadata thực thể. Đây là một triết lý huấn luyện rất đáng giữ lại vì nó phù hợp với đúng loại lỗi mà hệ retrieval tài chính thường mắc phải.",
            "Tuy nhiên, nếu đi sang bài toán reasoning toàn trình thì CACL theo dạng hiện tại vẫn chưa đủ. Lý do là lỗi của reasoning không còn chỉ nằm ở việc chọn sai tài liệu, mà nằm ở việc chọn sai toán hạng trong tập context đã truy xuất, chuẩn hóa sai đơn vị hoặc sinh sai chương trình tính toán. Khi đó, huấn luyện chỉ dựa vào contrastive retrieval sẽ không tác động trực tiếp vào phần quyết định nhất của bài toán.",
            "Vì vậy, chiến lược hợp lý hơn là xem CACL như nền retrieval-oriented, sau đó bổ sung một tầng huấn luyện mới cho reasoning. Tầng này nên kết hợp Supervised Fine-Tuning cho chương trình suy luận, preference optimization ở mức trace và reinforcement learning với reward có thể kiểm chứng bằng executor. Đây là lý do phần sau của báo cáo đề xuất pipeline SFT rồi Step-DPO rồi GRPO thay vì chỉ mở rộng thêm contrastive learning thuần retrieval.",
        ],
    )

    doc.add_heading("5.1. So sánh DPO, ORPO và GRPO cho bài toán hiện tại", level=2)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Để lựa chọn thuật toán huấn luyện phù hợp, cần đặt câu hỏi theo đặc thù bài toán tài chính chứ không theo độ phổ biến chung của từng thuật toán. Bảng dưới đây tổng hợp vai trò phù hợp của các lựa chọn chính.")

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = "Table Grid"
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers3 = ["Thuật toán", "Ưu điểm", "Hạn chế", "Vai trò đề xuất trong đề tài"]
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "F4CCCC")
    set_repeat_table_header(table3.rows[0])

    algo_rows = [
        ("DPO", "Ổn định, offline, dễ triển khai hơn RL online", "Tối ưu ở mức toàn câu trả lời, chưa khai thác mạnh reward thực thi", "Dùng sau SFT để chỉnh reasoning trace và tránh context nhiễu"),
        ("ORPO", "Chi phí thấp, phù hợp làm baseline", "Khó trở thành tối ưu chính cho numerical reasoning có executor", "Dùng như baseline so sánh compute thấp"),
        ("Step-DPO", "Nhấn mạnh chất lượng từng bước suy luận", "Cần chuẩn bị cặp trace chất lượng", "Lựa chọn ưu tiên hơn DPO thuần cho trace reasoning"),
        ("GRPO / RLVR", "Phù hợp bài toán có reward tính tự động từ executor", "Phụ thuộc mạnh vào chất lượng reward và verifier", "Thuật toán chính cho tối ưu reasoning cuối cùng"),
    ]
    for row in algo_rows:
        cells = table3.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    add_paragraphs(
        doc,
        [
            "Sau khi đối chiếu các thuật toán với bản chất bài toán tài chính, kết luận hợp lý nhất là không nên dùng ORPO như chiến lược chính và cũng không nên dừng ở DPO thuần. DPO có ích ở pha trung gian để dạy mô hình phân biệt trace grounded với trace dùng context nhiễu, nhưng nếu bài toán đã có executor và reward có thể kiểm chứng thì GRPO hoặc một biến thể RLVR mới là lựa chọn mạnh cho giai đoạn tối ưu cuối. Đồng thời, do lỗi reasoning thường nằm ở một bước trung gian như sai year hoặc sai unit, Step-DPO hoặc một biến thể process-aware sẽ phù hợp hơn response-level DPO.",
        ],
    )

    doc.add_heading("6. Chốt ý tưởng nghiên cứu và kiến trúc đề xuất", level=1)
    add_paragraphs(
        doc,
        [
            "Sau khi phân tích benchmark, code hiện tại và tài liệu liên quan, ý tưởng nên chốt không phải là một KG tốt hơn cho retrieval, mà là một Financial Evidence Graph dùng chung cho retrieval và reasoning. Đồ thị này cần liên kết được document, section, table, row, column, cell, sentence, footnote, company, period, metric concept, unit và equation. Chỉ khi đó KG mới thực sự trở thành cầu nối giữa truy xuất và suy luận, thay vì dừng ở vai trò structural reranker.",
            "Trong kiến trúc đề xuất, retrieval cần chuyển sang dạng phân cấp. Thay vì chỉ lấy top-k document rồi đưa nguyên văn cho mô hình suy luận, hệ sẽ đi qua ba tầng: document retrieval, table/section retrieval và evidence atom retrieval. Evidence atom ở đây có thể là cell, row aggregate, sentence hoặc footnote span. Từ các atom này, hệ mới xây local evidence graph cho từng câu hỏi cụ thể. Đây là khác biệt quyết định giữa một hệ retrieval-centric và một hệ evidence-grounded reasoning.",
            "Reasoning sau đó không chạy trên raw context mà chạy trên local evidence graph đã được triage. Câu hỏi sẽ được parser để xác định company, phạm vi thời gian, metric mục tiêu và loại phép toán. Module grounding sẽ chọn các toán hạng tương thích, sau đó một planner sẽ sinh DSL hoặc chương trình Python, executor thực thi phép tính và verifier kiểm tra ngược xem đáp án có nhất quán với company, year, đơn vị, scale và equation constraints hay không. Cách tiếp cận này biến bài toán trả lời câu hỏi thành bài toán dựng và kiểm chứng cấu trúc bằng chứng, phù hợp hơn với miền tài chính.",
        ],
    )

    doc.add_heading("6.1. Vì sao hướng này mạnh hơn việc chỉ cải thiện MRR", level=2)
    add_paragraphs(
        doc,
        [
            "Trong tài liệu tài chính, tăng MRR là cần nhưng chưa đủ. Một hệ có MRR@3 tốt nhưng thường xuyên trộn một số từ context đúng với một số từ context nhiễu vẫn không giải được bài toán thực. Điều advisor hoặc hội đồng đánh giá thực sự quan tâm là liệu hệ có đáng tin hay không khi đối diện tình huống retrieval không hoàn hảo. Kiến trúc dựa trên evidence graph và verifier trả lời đúng vào điểm đó, bởi nó không giả định rằng top-k context đã sạch mà chủ động giải quyết nhiễu trước và trong reasoning.",
            "Hướng mới cũng mạnh hơn về mặt câu chuyện học thuật. Nếu bài báo chỉ nói KG giúp retrieval tốt hơn thì rất dễ bị nhìn như một biến thể của graph reranking. Nhưng nếu bài báo cho thấy cùng một evidence graph được dùng để chọn bằng chứng, neo toán hạng, sinh chương trình và kiểm chứng kết quả thì contribution trở nên sâu hơn nhiều. Nó không còn là một mẹo retrieval mà là một cơ chế hợp nhất retrieval và reasoning trong miền tài chính.",
        ],
    )

    doc.add_heading("7. Kế hoạch triển khai chi tiết", level=1)
    add_paragraphs(
        doc,
        [
            "Lộ trình triển khai nên được chia thành các chặng rõ ràng để bảo đảm mỗi bước đều tạo ra kết quả trung gian có thể đánh giá và báo cáo. Việc này đặc biệt quan trọng vì bài toán toàn trình có độ rộng lớn; nếu không chia chặng hợp lý, hệ thống rất dễ rơi vào trạng thái nhiều ý tưởng nhưng không có mô-đun nào đủ chín.",
        ],
    )

    table4 = doc.add_table(rows=1, cols=4)
    table4.style = "Table Grid"
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers4 = ["Giai đoạn", "Mục tiêu", "Công việc chính", "Sản phẩm mong đợi"]
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, font_size=10)
        shade_cell(cell, "D9D2E9")
    set_repeat_table_header(table4.rows[0])

    roadmap_rows = [
        ("Giai đoạn 1", "Củng cố retrieval", "Mở rộng metadata schema, contextual chunk embedding, hard negatives mạnh hơn, table/section reranking", "Retrieval phân cấp mạnh hơn và có ablation rõ"),
        ("Giai đoạn 2", "Xây evidence graph", "Thiết kế schema node/edge, provenance, unit/scale, equation node, liên kết text-table-footnote", "Financial Evidence Graph hoạt động được"),
        ("Giai đoạn 3", "Xây reasoning substrate", "Query parser, operand grounding, DSL/Python executor, verifier", "Pipeline reasoning có thể thực thi và kiểm chứng"),
        ("Giai đoạn 4", "Huấn luyện reasoning", "SFT, Step-DPO, GRPO/RLVR với reward verifiable", "Mô hình reasoning ổn định và đáng tin hơn"),
        ("Giai đoạn 5", "Đánh giá toàn trình", "End-to-end benchmark, robustness dưới top-3 noisy contexts, case study", "Bộ kết quả hoàn chỉnh cho báo cáo và bài báo"),
    ]
    for row in roadmap_rows:
        cells = table4.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, font_size=9)

    add_paragraphs(
        doc,
        [
            "Thứ tự triển khai trong bảng trên không chỉ mang tính quản lý dự án mà còn phản ánh phụ thuộc kỹ thuật giữa các mô-đun. Retrieval phải đủ tốt thì evidence triage mới có đầu vào có ích; evidence graph phải rõ thì grounding và verifier mới đứng được; verifier phải đáng tin thì reward của RL mới không bị méo. Chính vì vậy, đề xuất huấn luyện SFT rồi Step-DPO rồi GRPO chỉ có ý nghĩa sau khi executor và verifier đã ổn định.",
        ],
    )

    doc.add_heading("8. Đóng góp khoa học dự kiến", level=1)
    add_paragraphs(
        doc,
        [
            "Nếu triển khai thành công, công trình có thể chốt ở ba đóng góp khoa học rõ ràng. Đóng góp thứ nhất là metadata-aware hierarchical retrieval cho tài liệu tài chính dài, trong đó metadata được nâng thành ontology truy xuất và retrieval được đẩy xuống cấp table/section rồi evidence atom. Đóng góp thứ hai là Financial Evidence Graph trung thành với phương trình, dùng chung cho retrieval và reasoning thay vì chỉ phục vụ structural reranking. Đóng góp thứ ba là pipeline reasoning có thể kiểm chứng, kết hợp grounding, executor, verifier và tối ưu theo hướng Step-DPO rồi GRPO.",
            "Ba đóng góp này có mối liên hệ tự nhiên với nhau. Retrieval phân cấp giải quyết việc đưa đúng bằng chứng vào hệ; evidence graph tạo không gian biểu diễn chung cho cả retrieval và reasoning; còn executor cùng verifier giúp chuyển từ RAG trả lời ngôn ngữ sang reasoning số học đáng tin. Chính sự gắn kết này làm cho đề tài có chiều sâu hơn một bài retrieval thuần hoặc một bài reasoning tách biệt khỏi truy xuất.",
        ],
    )

    doc.add_heading("9. Kết luận", level=1)
    add_paragraphs(
        doc,
        [
            "Báo cáo này đi từ hiện trạng triển khai thực tế trong repo, nội dung retrieval trong contribution1.pdf, các benchmark đang dùng, những khoảng trống kỹ thuật còn tồn tại, cho đến một chiến lược triển khai toàn trình đủ rõ để có thể báo cáo và làm việc tiếp theo với giáo viên hướng dẫn. Kết luận tổng quát là đề tài hiện đã có một prototype retrieval có cấu trúc đủ tốt để làm nền, nhưng chưa nên dừng ở đó.",
            "Hướng phát triển mạnh nhất là biến KG hiện tại thành Financial Evidence Graph dùng chung cho retrieval và reasoning, đồng thời nâng metadata thành ontology truy xuất và thêm một chuỗi reasoning có executor cùng verifier. Trên nền tảng đó, chiến lược huấn luyện phù hợp nhất là retrieval ranking ở giai đoạn đầu, SFT cho reasoning, Step-DPO để làm sạch trace và GRPO để tối ưu cuối cùng bằng reward verifiable. Đây là lộ trình vừa khả thi về mặt kỹ thuật, vừa có câu chuyện học thuật đủ rõ để phát triển thành kết quả nghiên cứu mạnh.",
        ],
    )

    doc.add_heading("Tài liệu tham khảo chính", level=1)
    refs = [
        "Rafailov et al., Direct Preference Optimization: Your Language Model is Secretly a Reward Model, 2023. https://arxiv.org/abs/2305.18290",
        "Hong et al., ORPO: Monolithic Preference Optimization without Reference Model, EMNLP 2024. https://aclanthology.org/2024.emnlp-main.626/",
        "Shao et al., DeepSeekMath: Pushing the Limits of Mathematical Reasoning in Open Language Models, 2024. https://arxiv.org/abs/2402.03300",
        "Chen et al., FinQA: A Dataset of Numerical Reasoning over Financial Data, EMNLP 2021. https://aclanthology.org/2021.emnlp-main.300/",
        "Zhu et al., TAT-QA: A Question Answering Benchmark on a Hybrid of Tabular and Textual Content in Finance, ACL 2021. https://aclanthology.org/2021.acl-long.254/",
        "Chen et al., ConvFinQA: Exploring the Chain of Numerical Reasoning in Conversational Finance Question Answering, EMNLP 2022. https://aclanthology.org/2022.emnlp-main.421/",
        "T2-RAGBench: Retrieval-first Benchmark for Text-and-Table Financial QA, 2025 preprint. https://arxiv.org/abs/2506.12071",
        "HierFinRAG: Hierarchical Financial Retrieval-Augmented Generation, 2026. https://www.mdpi.com/2227-9709/13/2/30",
        "FT-RAG: Fine-Grained Table Retrieval-Augmented Generation for Financial Analysis, 2026 preprint. https://arxiv.org/abs/2605.01495",
        "Metadata-Driven Financial RAG: Retrieval-Enhanced Large Language Models for Financial Question Answering, 2025 preprint. https://arxiv.org/abs/2510.24402",
        "Structure First, Reason Next: A Framework for Financial Numerical Reasoning with Structured Evidence, 2026 preprint. https://arxiv.org/abs/2601.07754",
        "Table-R1: Reinforcement Learning for Table Reasoning, 2025 preprint. https://arxiv.org/abs/2505.23621",
        "Contribution nội bộ: Structured Knowledge-Enhanced Retrieval for Financial Documents, contribution1.pdf trong thư mục NLP.",
    ]
    add_reference_list(doc, refs)

    doc.save(docx_path)


def main() -> None:
    build_diagram(FIG_PATH)
    build_report(OUT_DOCX, FIG_PATH)
    print(f"Generated: {OUT_DOCX}")
    print(f"Generated: {FIG_PATH}")


if __name__ == "__main__":
    main()
